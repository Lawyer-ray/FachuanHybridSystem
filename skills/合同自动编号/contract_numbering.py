#!/usr/bin/env python3
"""
Contract Numbering Skill
将合同文档中的手动编号转换为 Word 自动编号

支持两种格式：
1. 一、1.（1）① （中文格式）
2. 1. 1.1 1.1.1 1.1.1.1 1.1.1.1.1 （纯数字格式）
"""

import logging
import re
import sys
from pathlib import Path

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)

# 签字盖章部分的关键词
SIGNATURE_KEYWORDS = [
    '以下无正文',
    '签约页',
    '（盖章）',
    '(盖章)',
    '授权代表签字',
    '签订日期',
    '甲方（盖章）',
    '乙方（盖章）',
    '丙方（盖章）',
    '丁方（盖章）',
]

# 编号格式定义
NUMBERING_FORMATS = {
    'chinese': {
        'name': '一、1.（1）①',
        'description': '中文格式（一、二、三...）',
        'levels': [
            {'ilvl': '0', 'numFmt': 'chineseCounting', 'lvlText': '%1、', 'start': '1'},
            {'ilvl': '1', 'numFmt': 'decimal', 'lvlText': '%2.', 'start': '1'},
            {'ilvl': '2', 'numFmt': 'decimal', 'lvlText': '（%3）', 'start': '1'},
            {'ilvl': '3', 'numFmt': 'decimalEnclosedCircle', 'lvlText': '%4', 'start': '1'},
        ]
    },
    'decimal': {
        'name': '1. 1.1 1.1.1 1.1.1.1',
        'description': '纯数字格式（1. 2. 3...）',
        'levels': [
            {'ilvl': '0', 'numFmt': 'decimal', 'lvlText': '%1.', 'start': '1'},
            {'ilvl': '1', 'numFmt': 'decimal', 'lvlText': '%1.%2.', 'start': '1'},
            {'ilvl': '2', 'numFmt': 'decimal', 'lvlText': '%1.%2.%3.', 'start': '1'},
            {'ilvl': '3', 'numFmt': 'decimal', 'lvlText': '%1.%2.%3.%4.', 'start': '1'},
            {'ilvl': '4', 'numFmt': 'decimal', 'lvlText': '%1.%2.%3.%4.%5.', 'start': '1'},
        ]
    }
}


def is_signature_section(text: str) -> bool:
    """检测是否是签字盖章部分"""
    for keyword in SIGNATURE_KEYWORDS:
        if keyword in text:
            return True
    return False


def detect_numbering_structure(doc: Document, format_type: str = 'chinese') -> list:
    """
    AI 分析文档结构，识别编号层级

    Args:
        doc: Word 文档
        format_type: 编号格式类型 ('chinese' 或 'decimal')

    Returns:
        list of (para_idx, level, matched_text, original_text)
    """
    # 识别 Level 0 标题
    level0_paras = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        if format_type == 'chinese':
            # 中文格式：一、二、三...
            m = re.match(r'^([一二三四五六七八九十]+)、\s*', text)
            if m:
                level0_paras.append((i, m.group(0), text))
        else:
            # 纯数字格式：1. 2. 3...（只匹配顶级数字）
            m = re.match(r'^(\d+)\.\s+', text)
            if m and '.' not in text[len(m.group(0)):len(m.group(0))+5]:
                level0_paras.append((i, m.group(0), text))

    # 识别所有编号段落
    numbered_paras = []
    paras = doc.paragraphs

    for level0_idx, (para_idx, matched, text) in enumerate(level0_paras):
        # Level 0 标题
        numbered_paras.append((para_idx, 0, matched, text))

        # 找出下一个 Level 0 的位置
        next_level0_idx = level0_paras[level0_idx + 1][0] if level0_idx + 1 < len(level0_paras) else len(paras)

        # 扫描 Level 0 下的所有段落
        has_level1_heading = False
        prev_level = 0

        for i in range(para_idx + 1, next_level0_idx):
            para = paras[i]
            text = para.text.strip()
            if not text:
                continue

            # 检测签字盖章部分，停止编号
            if is_signature_section(text):
                break

            if format_type == 'chinese':
                # 中文格式检测逻辑

                # Level 1: （一）（二）... 或 （1）（2）...
                m = re.match(r'^[（(]([一二三四五六七八九十\d]+)[）)]\s*', text)
                if m:
                    numbered_paras.append((i, 1, m.group(0), text))
                    has_level1_heading = True
                    prev_level = 1
                    continue

                # 数字编号：1. 2. 3... 或 1.1 1.2...
                m = re.match(r'^(\d+\.\d+|\d+[.、])\s*', text)
                if m:
                    if has_level1_heading:
                        numbered_paras.append((i, 2, m.group(0), text))
                        prev_level = 2
                    else:
                        numbered_paras.append((i, 1, m.group(0), text))
                        prev_level = 1
                    continue

            else:
                # 纯数字格式检测逻辑

                # 检测数字层级：1.1.1.1.1 > 1.1.1.1 > 1.1.1 > 1.1 > 1.
                m = re.match(r'^(\d+(?:\.\d+){0,4})\.\s+', text)
                if m:
                    num_str = m.group(1)
                    level = num_str.count('.')  # 点号数量决定层级
                    if level > 4:
                        level = 4  # 最多5级

                    numbered_paras.append((i, level, m.group(0), text))
                    prev_level = level
                    continue

            # 其他段落：继承上一个段落的级别
            if prev_level >= 1:
                numbered_paras.append((i, prev_level, '', text))
            else:
                numbered_paras.append((i, 1, '', text))

    return numbered_paras


def create_numbering_part(doc: Document):
    """创建或获取 numbering part"""
    try:
        return doc.part.numbering_part, doc.part.numbering_part.element
    except:
        package = doc.part.package
        partnames = {part.partname for part in package.iter_parts()}
        partname = PackURI('/word/numbering.xml')
        if partname in partnames:
            partname = package.next_partname('/word/numbering%d.xml')

        numbering_xml = '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        numbering_part = Part(
            partname,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
            numbering_xml.encode('utf-8'),
            package
        )
        doc.part.relate_to(numbering_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
        numbering_elem = etree.fromstring(numbering_xml.encode('utf-8'))
        return numbering_part, numbering_elem


def create_abstract_numbering(numbering_elem, abstract_id: int, format_type: str = 'chinese') -> None:
    """创建 abstractNum 定义"""
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_id))

    multi_level = OxmlElement('w:multiLevelType')
    multi_level.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level)

    # 获取格式定义
    format_def = NUMBERING_FORMATS.get(format_type, NUMBERING_FORMATS['chinese'])
    levels = format_def['levels']

    for lvl_def in levels:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), lvl_def['ilvl'])

        start = OxmlElement('w:start')
        start.set(qn('w:val'), lvl_def['start'])
        lvl.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), lvl_def['numFmt'])
        lvl.append(num_fmt)

        suff = OxmlElement('w:suff')
        suff.set(qn('w:val'), 'nothing')
        lvl.append(suff)

        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), lvl_def['lvlText'])
        lvl.append(lvl_text)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:leftChars'), '0')
        ind.set(qn('w:hanging'), '0')
        ind.set(qn('w:hangingChars'), '0')
        p_pr.append(ind)
        lvl.append(p_pr)

        abstract_num.append(lvl)

    numbering_elem.append(abstract_num)


def create_num_instances(numbering_elem, abstract_id: int, level0_indices: list, format_type: str = 'chinese') -> dict:
    """为每个 Level 0 创建独立的 num 实例"""
    num_id_map = {}
    next_num_id = 1

    # 获取格式定义中的最大层级
    format_def = NUMBERING_FORMATS.get(format_type, NUMBERING_FORMATS['chinese'])
    max_level = len(format_def['levels']) - 1

    for para_idx in level0_indices:
        num_elem = OxmlElement('w:num')
        num_elem.set(qn('w:numId'), str(next_num_id))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(abstract_id))
        num_elem.append(abstract_ref)

        # 重置 Level 1 到 max_level 的计数器
        for reset_level in range(1, max_level + 1):
            lvl_override = OxmlElement('w:lvlOverride')
            lvl_override.set(qn('w:ilvl'), str(reset_level))
            start_override = OxmlElement('w:startOverride')
            start_override.set(qn('w:val'), '1')
            lvl_override.append(start_override)
            num_elem.append(lvl_override)

        numbering_elem.append(num_elem)
        num_id_map[para_idx] = next_num_id
        next_num_id += 1

    return num_id_map


def apply_numbering(doc: Document, numbered_paras: list, num_id_map: dict) -> None:
    """应用自动编号到段落"""
    current_num_id = None

    for para_idx, level, matched, original_text in numbered_paras:
        para = doc.paragraphs[para_idx]

        # 去除手动编号
        if matched:
            new_text = original_text[len(matched):].lstrip()
        else:
            new_text = original_text

        # 清除并重设文本
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text

        # 确定 numId
        if level == 0:
            current_num_id = num_id_map.get(para_idx)

        # 应用自动编号
        p_pr = para._element.find(qn('w:pPr'))
        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            para._element.insert(0, p_pr)

        # 移除旧的 numPr
        old_num_pr = p_pr.find(qn('w:numPr'))
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)

        # 移除旧的缩进
        old_ind = p_pr.find(qn('w:ind'))
        if old_ind is not None:
            p_pr.remove(old_ind)

        # 创建新的 numPr
        num_pr = OxmlElement('w:numPr')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        num_pr.append(ilvl)

        num_id_elem = OxmlElement('w:numId')
        num_id_to_use = current_num_id if level >= 1 else num_id_map.get(para_idx)
        num_id_elem.set(qn('w:val'), str(num_id_to_use))
        num_pr.append(num_id_elem)

        p_pr.append(num_pr)

        # 设置缩进为 0
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:leftChars'), '0')
        ind.set(qn('w:hanging'), '0')
        ind.set(qn('w:hangingChars'), '0')
        p_pr.append(ind)


def get_user_format_choice() -> str:
    """询问用户选择编号格式"""
    logger.info("\n请选择编号格式：")
    logger.info("  1. 一、1.（1）① （中文格式）")
    logger.info("  2. 1. 1.1 1.1.1 1.1.1.1 （纯数字格式）")
    logger.info("")

    while True:
        choice = input("请输入选项 (1 或 2): ").strip()
        if choice == '1':
            return 'chinese'
        elif choice == '2':
            return 'decimal'
        else:
            logger.warning("无效选项，请输入 1 或 2")


def convert_contract_numbering(input_path: str, output_path: str = None, format_type: str = None) -> dict:
    """
    转换合同文档的自动编号

    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径（可选，默认为 {原文件名}_自动编号.docx）
        format_type: 编号格式类型 ('chinese' 或 'decimal')，None 则询问用户

    Returns:
        dict: 转换结果信息
    """
    input_path = Path(input_path)
    if not input_path.exists():
        return {'success': False, 'error': f'文件不存在: {input_path}'}

    if output_path is None:
        output_path = input_path.parent / f"{input_path.stem}_自动编号{input_path.suffix}"

    # 如果未指定格式，询问用户
    if format_type is None:
        format_type = get_user_format_choice()

    # 读取文档
    doc = Document(input_path)

    # 分析文档结构
    numbered_paras = detect_numbering_structure(doc, format_type)

    # 提取 Level 0 索引
    level0_indices = [idx for idx, level, _, _ in numbered_paras if level == 0]

    # 创建 numbering part
    numbering_part, numbering_elem = create_numbering_part(doc)

    # 创建 abstractNum
    create_abstract_numbering(numbering_elem, abstract_id=0, format_type=format_type)

    # 创建 num 实例
    num_id_map = create_num_instances(numbering_elem, abstract_id=0, level0_indices=level0_indices, format_type=format_type)

    # 更新 numbering part
    if hasattr(numbering_part, '_blob'):
        numbering_part._blob = etree.tostring(numbering_elem, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 应用自动编号
    apply_numbering(doc, numbered_paras, num_id_map)

    # 保存文档
    doc.save(output_path)

    return {
        'success': True,
        'input_path': str(input_path),
        'output_path': str(output_path),
        'format_type': format_type,
        'format_name': NUMBERING_FORMATS[format_type]['name'],
        'total_paragraphs': len(numbered_paras),
        'level0_count': len(level0_indices),
        'numbered_paras': numbered_paras,
    }


def main():
    """命令行入口"""
    logging.basicConfig(level=logging.INFO, format='%(message)s')

    if len(sys.argv) < 2:
        logger.error("用法: python contract_numbering.py <input_docx> [output_docx] [--format chinese|decimal]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = None
    format_type = None

    # 解析命令行参数
    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == '--format' and i + 1 < len(sys.argv):
            format_type = sys.argv[i + 1]
            i += 2
        elif output_path is None:
            output_path = sys.argv[i]
            i += 1
        else:
            i += 1

    result = convert_contract_numbering(input_path, output_path, format_type)

    if result['success']:
        logger.info("✓ 转换成功")
        logger.info("  输入: %s", result['input_path'])
        logger.info("  输出: %s", result['output_path'])
        logger.info("  格式: %s", result['format_name'])
        logger.info("  总段落数: %d", result['total_paragraphs'])
        logger.info("  一级标题数: %d", result['level0_count'])

        # 显示转换映射
        logger.info("\n=== 转换映射 ===\n")
        for idx, level, matched, text in result['numbered_paras'][:30]:
            if level == 0:
                logger.info("\n--- %s ---", text[:20])
            else:
                prefix = '  ' * level
                note = ' [推断]' if matched == '' else ''
                logger.info("[%3d] %sL%d: %s%s", idx, prefix, level, text[:40], note)

        if len(result['numbered_paras']) > 30:
            logger.info("\n... 还有 %d 个段落", len(result['numbered_paras']) - 30)
    else:
        logger.error("✗ 转换失败: %s", result['error'])
        sys.exit(1)


if __name__ == '__main__':
    main()
