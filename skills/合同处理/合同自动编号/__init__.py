"""
合同自动编号 Skill

将合同文档中的手动编号转换为 Word 自动编号。

支持两种格式：
1. 一、1.（1）① （中文格式）
2. 1. 1.1 1.1.1 1.1.1.1 1.1.1.1.1 （纯数字格式）

支持两种工作模式：
- 自动模式：正则检测编号结构（适合格式规范的合同）
- AI 辅助模式：AI 判断前缀和层级，运行时直接使用（适合任何格式）
"""

import json
import logging
import re
from pathlib import Path

from docx import Document

from .converter import convert_numbering, verify_numbering
from .detector import detect_chinese_level0, detect_decimal_level0, detect_numbering_structure, is_signature_section
from .formats import NUMBERING_FORMATS, get_format
from .utils import format_numbering_mapping, generate_output_path, validate_input_path

logger = logging.getLogger(__name__)

__version__ = '1.6.0'
__all__ = ['convert_contract_numbering', 'analyze_document', 'apply_numbering_map', 'NUMBERING_FORMATS']

# 最大重试次数
MAX_RETRIES = 3


def _extract_prefix(text: str) -> str:
    """提取段落开头的编号前缀（自动模式辅助用）

    注意：AI 辅助模式下，前缀已由 AI 直接提供，不再使用此函数。
    此函数仅用于自动模式，支持常见格式（一、1.（1）等）。
    """
    patterns = [
        r'^第[一二三四五六七八九十0-9]+[条项段款](?:\s*[：:]?\s+|\s+)',
        r'^\d+[-\.·]\d+(?:[-\.·]\d+){0,3}(?:[、。！？：；\s．.])*',
        r'^[（(][一二三四五六七八九十]+[）)]\s*',
        r'^[（(]\d+[）)]\s*',
        r'^[一二三四五六七八九十]+、\s*',
        r'^\d+\.\d+(?:\.\d+){0,3}[.、．.]?\s*',
        r'^\d+[.、．.]\s*',
        r'^\d+[）)]\s*',
    ]
    for pattern in patterns:
        m = re.match(pattern, text)
        if m:
            return m.group(0)
    return ''


def get_user_format_choice() -> str:
    """询问用户选择编号格式

    Returns:
        格式类型 ('chinese' 或 'decimal')
    """
    logger.info("\n请选择编号格式：")
    logger.info("  1. 一、1.（1）① （中文格式）")
    logger.info("  2. 1. 1.1 1.1.1 1.1.1.1 （纯数字格式）")
    logger.info("")

    while True:
        choice = input("请输入选项 (1 或 2): ").strip()
        if choice == '1':
            return 'chinese'
        elif choice == '2':
            return 'decimal'
        else:
            logger.warning("无效选项，请输入 1 或 2")


def convert_contract_numbering(
    input_path: str | Path,
    output_path: str | Path | None = None,
    format_type: str | None = None,
    verify: bool = True
) -> dict:
    """
    转换合同文档的自动编号

    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径（可选，默认为 {原文件名}_自动编号.docx）
        format_type: 编号格式类型 ('chinese' 或 'decimal')，None 则询问用户
        verify: 是否验证转换结果

    Returns:
        dict: 转换结果信息（含审计报告）
    """
    # 验证输入路径
    try:
        input_path = validate_input_path(input_path)
    except (FileNotFoundError, ValueError) as e:
        return {'success': False, 'error': str(e)}

    # 生成输出路径
    if output_path is None:
        output_path = generate_output_path(input_path)
    else:
        output_path = Path(output_path)

    # 如果未指定格式，询问用户
    if format_type is None:
        format_type = get_user_format_choice()

    # 验证格式类型
    try:
        get_format(format_type)
    except ValueError as e:
        return {'success': False, 'error': str(e)}

    # 预读取原始文档（供审计使用）
    original_doc_for_audit = Document(input_path)

    # 重试循环
    for attempt in range(1, MAX_RETRIES + 1):
        logger.info("尝试第 %d 次转换...", attempt)

        # 读取文档（每次重新读取以重新开始）
        doc = Document(input_path)

        # 分析文档结构
        numbered_paras = detect_numbering_structure(doc, format_type)

        # 提取 Level 0 索引
        level0_indices = [idx for idx, level, _, _ in numbered_paras if level == 0]

        # 转换编号
        convert_numbering(doc, numbered_paras, {}, format_type)

        # 验证转换结果
        if verify:
            verification = verify_numbering(
                doc, numbered_paras,
                original_doc=original_doc_for_audit,
                format_type=format_type,
            )

            if verification['all_valid']:
                logger.info("✓ 验证通过！所有 %d 个段落编号正确", verification['total'])

                # 保存文档
                doc.save(output_path)

                return {
                    'success': True,
                    'input_path': str(input_path),
                    'output_path': str(output_path),
                    'format_type': format_type,
                    'format_name': NUMBERING_FORMATS[format_type]['name'],
                    'total_paragraphs': len(numbered_paras),
                    'level0_count': len(level0_indices),
                    'numbered_paras': numbered_paras,
                    'verification': verification,
                    'audit': verification.get('audit'),
                }
            else:
                logger.warning("✗ 验证失败！%d/%d 个段落编号不正确",
                             verification['invalid_count'], verification['total'])

                # 显示失败详情
                for r in verification['results']:
                    if not r['valid']:
                        expected = f"L{r['expected_level']}" if r['expected_level'] is not None else "无编号"
                        actual = f"L{r['actual_level']}" if r['actual_level'] is not None else "无"
                        logger.warning("  [%d] 期望 %s, 实际 %s: %s",
                                     r['para_idx'], expected, actual, r['text'])

                # 显示审计报告（增强信息）
                if verification.get('audit'):
                    audit = verification['audit']
                    if not audit['all_clear']:
                        logger.warning("\n⚠ 审计发现以下问题：")
                        logger.warning(audit['summary'])

                if attempt < MAX_RETRIES:
                    logger.info("正在重试...")
                else:
                    logger.error("已达到最大重试次数 %d，转换失败", MAX_RETRIES)
                    return {
                        'success': False,
                        'error': f'验证失败：{verification["invalid_count"]} 个段落编号不正确',
                        'verification': verification,
                        'audit': verification.get('audit'),
                    }
        else:
            # 不验证，直接保存
            doc.save(output_path)

            return {
                'success': True,
                'input_path': str(input_path),
                'output_path': str(output_path),
                'format_type': format_type,
                'format_name': NUMBERING_FORMATS[format_type]['name'],
                'total_paragraphs': len(numbered_paras),
                'level0_count': len(level0_indices),
                'numbered_paras': numbered_paras,
            }

    # 不应该到这里
    return {'success': False, 'error': '未知错误'}


# ============================================================
# AI 辅助模式
# ============================================================


def _analyze_prefix(text: str) -> tuple[str, str]:
    """分析文本中的编号前缀，返回 (编号前缀, 去除编号后的正文)

    使用分段扫描策略，逐步识别段落开头的编号部分。
    适用于任何格式，包括非标准编号如 "1-1、" "第一条" 等。
    """
    text = text.lstrip()
    if not text:
        return '', ''

    # 前缀候选模式（按优先级从高到低）
    prefix_patterns = [
        # 「第」+ 汉字/数字 + 量词 + 标点/空格
        r'^第[一二三四五六七八九十〇0-9]+(?:[条项款段章])+[、。：；\s]*',
        # 连字符多级 1-1、 4-1-2. 3·1·4（允许末尾无标号）
        r'^\d+[-·－]\d+(?:[-·－]\d+){0,3}(?:[、。：；．.，,（(\)\s])*',
        # 括号中文（一）（二）
        r'^[（(][一二三四五六七八九十]+[）)][、。：；\s]*',
        # 括号数字（1）（2）
        r'^[（(]\d+[）)][、。：；\s]*',
        # 汉字顿号 一、 二、
        r'^[一二三四五六七八九十]+、',
        # 多级点号 1.1.1. 1.1.1.1
        r'^\d+\.\d+(?:\.\d+){0,4}[.、。：；\s]*',
        # 单级数字 1. 2、 3．
        r'^\d+[.、．]',
        # 中文圈号 ①②③
        r'^[①②③④⑤⑥⑦⑧⑨⑩]+',
        # 字母编号 a) b)
        r'^[a-zA-Z][）).．]',
    ]

    for pattern in prefix_patterns:
        m = re.match(pattern, text)
        if m:
            prefix = m.group(0)
            body = text[len(prefix):].lstrip()
            return prefix, body

    return '', text


def analyze_document(input_path: str | Path, format_type: str = 'chinese') -> str:
    """分析文档段落结构，输出 JSON 供 AI 判断层级和前缀

    流程：
    1. 读取文档每个段落
    2. 用 _analyze_prefix 提取文号前缀和正文
    3. 输出结构化 JSON
    4. AI 读取后，判断每个段落的层级并填入 prefix 字段
    5. 生成 numbering_map.json

    编号层级定义：
    - level: -1 = 不编号（前后文/签名区）
    - level: 0 = 一级标题（如"第一条""一、"）
    - level: 1 = 二级标题（如"1-1、""1."）
    - level: 2 = 三级标题（如"4-1-1、""（1）"）
    - level: 3 = 四级标题（如"①"）

    Args:
        input_path: 输入文档路径
        format_type: 编号格式类型（影响输出格式建议）

    Returns:
        JSON 字符串，含 "format_type" 和 "paragraphs" 列表
    """
    input_path = validate_input_path(input_path)
    get_format(format_type)  # 验证格式
    doc = Document(input_path)

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        prefix, body = _analyze_prefix(text)

        # 检测是否是签名区
        is_sig = is_signature_section(text)
        # 签名区以外的空 prefix + 短文本段落通常是不编号的前言/正文
        is_plain = not is_sig and not prefix and len(text) > 20

        # 初步判断是否是 Level 0（AI 可以覆盖这个判断）
        if prefix and format_type == 'chinese':
            matched, _ = detect_chinese_level0(text)
            is_level0 = bool(matched)
        elif prefix:
            matched, _ = detect_decimal_level0(text)
            is_level0 = bool(matched)
        else:
            is_level0 = False

        paragraphs.append({
            'index': i,
            'text': text[:150],
            'prefix': prefix,
            'prefix_hint': 'AI 请修正：把编号部分（如"第一条 ""1-1、""4-1-1、"）完整填入',
            'is_level0': is_level0,
            'is_signature': is_sig,
            'is_plain': is_plain,
            'suggested_level': 0 if is_level0 else (-1 if is_sig else None),
        })

    result = {
        'format_type': format_type,
        'total_paragraphs': len(doc.paragraphs),
        'note': 'AI 辅助模式：请读取每个段落的 prefix 字段，确认或修正层级后输出 numbering_map.json',
        'numbering_map_format': [
            {'index': '段落号', 'level': '层级', 'prefix': 'AI 填写的编号前缀（必须完整）'}
        ],
        'paragraphs': paragraphs,
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def apply_numbering_map(
    input_path: str | Path,
    numbering_map: str | list[dict],
    output_path: str | Path | None = None,
    format_type: str = 'chinese',
    verify: bool = True,
) -> dict:
    """根据 AI 提供的层级映射应用自动编号

    关键改进：
    - numbering_map 中每个条目必须有 "index"、"level" 和 "prefix"
    - "prefix" 由 AI 填写，运行时直接用于切除旧编号
    - 无需运行时正则猜测编号前缀

    Args:
        input_path: 输入文档路径
        numbering_map: AI 生成的层级映射，格式：
            [{"index": 10, "level": 0, "prefix": "第一条 "},
             {"index": 26, "level": 1, "prefix": "1-1、"},
             ...]
            level: -1=不编号 0=一级 1=二级 2=三级 3=四级
            prefix: 必须完整，AI 直接判断的编号前缀
        output_path: 输出文档路径
        format_type: 编号格式类型
        verify: 是否验证转换结果

    Returns:
        dict: 转换结果信息
    """
    # 验证输入路径
    try:
        input_path = validate_input_path(input_path)
    except (FileNotFoundError, ValueError) as e:
        return {'success': False, 'error': str(e)}

    # 生成输出路径
    if output_path is None:
        output_path = generate_output_path(input_path)
    else:
        output_path = Path(output_path)

    # 验证格式
    try:
        get_format(format_type)
    except ValueError as e:
        return {'success': False, 'error': str(e)}

    # 解析 numbering_map
    if isinstance(numbering_map, str):
        mapping_data = json.loads(numbering_map)
    else:
        mapping_data = numbering_map

    # 读取文档
    original_doc_for_audit = Document(input_path)
    doc = Document(input_path)

    # ✅ 关键改进：使用 AI 提供的 prefix，直接切除旧编号
    numbered_paras = []
    missing_prefix_items = []
    for item in mapping_data:
        idx = item['index']
        level = item['level']
        para = doc.paragraphs[idx]
        text = para.text.strip()

        # 获取 prefix：优先用 AI 提供的，否则回退（只允许 level>=0 的段落有 prefix）
        prefix = item.get('prefix', '')
        if not prefix and level >= 0:
            # AI 没填 prefix，但指定了编号层级 → 尝试自动提取（作为兼容回退）
            missing_prefix_items.append(item)
            prefix = _extract_prefix(text)

        numbered_paras.append((idx, level, prefix, text))

    if missing_prefix_items:
        logger.warning(
            "警告：%d 个编号段落的 prefix 为空，已使用自动提取结果。"
            "建议在 numbering_map.json 中填写准确的 prefix，\n"
            "参考样式：{\"index\": 26, \"level\": 1, \"prefix\": \"1-1、\"}",
            len(missing_prefix_items)
        )

    # 只处理需要编号的段落（level >= 0）
    numbered_paras = [(idx, lvl, pfx, txt) for idx, lvl, pfx, txt in numbered_paras if lvl >= 0]

    # 提取 Level 0 索引
    level0_indices = [idx for idx, level, _, _ in numbered_paras if level == 0]

    if not level0_indices:
        return {'success': False, 'error': '层级映射中没有 Level 0 段落，无法应用编号'}

    # 应用编号
    convert_numbering(doc, numbered_paras, {}, format_type)

    # 验证
    verification = None
    if verify:
        verification = verify_numbering(
            doc, numbered_paras,
            original_doc=original_doc_for_audit,
            format_type=format_type,
        )

        if not verification['all_valid']:
            logger.warning("验证发现问题：\n%s", verification.get('audit', {}).get('summary', ''))

    # 保存
    doc.save(output_path)

    return {
        'success': True,
        'input_path': str(input_path),
        'output_path': str(output_path),
        'format_type': format_type,
        'format_name': NUMBERING_FORMATS[format_type]['name'],
        'total_paragraphs': len(numbered_paras),
        'level0_count': len(level0_indices),
        'numbered_paras': numbered_paras,
        'verification': verification,
    }
