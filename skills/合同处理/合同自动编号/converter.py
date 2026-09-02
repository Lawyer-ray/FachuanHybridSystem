"""
编号转换模块

将检测到的手动编号转换为 Word 自动编号。
"""

from __future__ import annotations

import logging

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from .formats import get_format, get_max_level

logger = logging.getLogger(__name__)


def create_numbering_part(doc: Document) -> tuple[Part, etree._Element]:
    """创建或获取 numbering part

    Args:
        doc: Word 文档

    Returns:
        (numbering_part, numbering_elem) 元组
    """
    try:
        return doc.part.numbering_part, doc.part.numbering_part.element
    except (AttributeError, Exception):
        package = doc.part.package
        partnames = {part.partname for part in package.iter_parts()}
        partname = PackURI('/word/numbering.xml')
        if partname in partnames:
            partname = package.next_partname('/word/numbering%d.xml')

        numbering_xml = '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        numbering_part = Part(
            partname,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
            numbering_xml.encode('utf-8'),
            package
        )
        doc.part.relate_to(numbering_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
        numbering_elem = etree.fromstring(numbering_xml.encode('utf-8'))
        return numbering_part, numbering_elem


def next_abstract_id(numbering_elem: etree._Element) -> int:
    """计算可用的下一个 abstractNum id，避免与文档已有定义冲突

    原实现硬编码 abstract_id=0，但源文档 numbering.xml 常已存在 id=0 的
    旧 abstractNum，新增的同 id 定义会被忽略，导致编号样式错乱。
    """
    ids = [
        int(a.get(qn('w:abstractNumId'), 0)) or 0
        for a in numbering_elem.findall(qn('w:abstractNum'))
    ]
    return max(ids, default=-1) + 1


def create_abstract_numbering(numbering_elem: etree._Element, abstract_id: int, format_type: str) -> None:
    """创建 abstractNum 定义

    Args:
        numbering_elem: numbering XML 元素
        abstract_id: abstractNum ID
        format_type: 格式类型
    """
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_id))

    multi_level = OxmlElement('w:multiLevelType')
    multi_level.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level)

    # 获取格式定义
    format_def = get_format(format_type)
    levels = format_def['levels']

    for lvl_def in levels:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), lvl_def['ilvl'])

        start = OxmlElement('w:start')
        start.set(qn('w:val'), lvl_def['start'])
        lvl.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), lvl_def['numFmt'])
        lvl.append(num_fmt)

        suff = OxmlElement('w:suff')
        suff.set(qn('w:val'), 'nothing')
        lvl.append(suff)

        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), lvl_def['lvlText'])
        lvl.append(lvl_text)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:leftChars'), '0')
        ind.set(qn('w:hanging'), '0')
        ind.set(qn('w:hangingChars'), '0')
        p_pr.append(ind)
        lvl.append(p_pr)

        abstract_num.append(lvl)

    numbering_elem.append(abstract_num)


def create_num_instances(
    numbering_elem: etree._Element,
    abstract_id: int,
    level0_indices: list[int],
    format_type: str
) -> dict[int, int]:
    """为每个 Level 0 创建独立的 num 实例

    Args:
        numbering_elem: numbering XML 元素
        abstract_id: abstractNum ID
        level0_indices: Level 0 段落索引列表
        format_type: 格式类型

    Returns:
        {para_idx: num_id} 映射
    """
    num_id_map = {}

    # 找出已存在的最大 numId，避免冲突
    existing_num_ids = [
        int(num.get(qn('w:numId'), 0))
        for num in numbering_elem.findall(qn('w:num'))
    ]
    next_num_id = max(existing_num_ids, default=0) + 1

    # 获取格式定义中的最大层级
    max_level = get_max_level(format_type)

    for para_idx in level0_indices:
        num_elem = OxmlElement('w:num')
        num_elem.set(qn('w:numId'), str(next_num_id))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(abstract_id))
        num_elem.append(abstract_ref)

        # 重置 Level 1 到 max_level 的计数器
        for reset_level in range(1, max_level + 1):
            lvl_override = OxmlElement('w:lvlOverride')
            lvl_override.set(qn('w:ilvl'), str(reset_level))
            start_override = OxmlElement('w:startOverride')
            start_override.set(qn('w:val'), '1')
            lvl_override.append(start_override)
            num_elem.append(lvl_override)

        numbering_elem.append(num_elem)
        num_id_map[para_idx] = next_num_id
        next_num_id += 1

    return num_id_map


def apply_numbering_to_paragraph(para, level: int, num_id: int, text: str) -> None:
    """应用自动编号到单个段落

    Args:
        para: 段落元素
        level: 编号层级
        num_id: 编号 ID
        text: 段落文本（已去除手动编号）
    """
    # 清除并重设文本
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text

    # 应用自动编号
    p_pr = para._element.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = OxmlElement('w:pPr')
        para._element.insert(0, p_pr)

    # 移除旧的 numPr
    old_num_pr = p_pr.find(qn('w:numPr'))
    if old_num_pr is not None:
        p_pr.remove(old_num_pr)

    # 移除旧的缩进
    old_ind = p_pr.find(qn('w:ind'))
    if old_ind is not None:
        p_pr.remove(old_ind)

    # 创建新的 numPr
    num_pr = OxmlElement('w:numPr')

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    num_pr.append(ilvl)

    num_id_elem = OxmlElement('w:numId')
    num_id_elem.set(qn('w:val'), str(num_id))
    num_pr.append(num_id_elem)

    p_pr.append(num_pr)

    # 设置缩进为 0
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:hanging'), '0')
    ind.set(qn('w:hangingChars'), '0')
    p_pr.append(ind)


def _remove_numPr(para) -> bool:
    """清除段落的 numPr 列表编号属性，返回是否清除了"""
    p_pr = para._element.find(qn('w:pPr'))
    if p_pr is not None:
        old_num_pr = p_pr.find(qn('w:numPr'))
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)
            return True
    return False


def _is_heading_style(para, heading_styles: set[str]) -> bool:
    """判断段落是否为 Heading 样式"""
    try:
        style_name = para.style.name if para.style else ''
        return style_name in heading_styles
    except (AttributeError, KeyError):
        return False


def _should_skip_paragraph(text: str) -> bool:
    """判断段落是否应该跳过（纯标题行/不应该编号）"""
    if not text.strip():
        return True

    # 纯标题行：短文本，无冒号/说明内容，如"甲方信息""乙方/实控人信息"
    if len(text.strip()) < 15 and ':' not in text and '：' not in text and '。' not in text:
        return True

    return False


def convert_numbering(
    doc: Document,
    numbered_paras: list[tuple[int, int, str, str]],
    num_id_map: dict[int, int],
    format_type: str
) -> None:
    """转换文档中的编号

    Args:
        doc: Word 文档
        numbered_paras: 编号段落列表，格式 (para_idx, level, clean_text, original_text)
            clean_text: 段落正文（不含编号前缀），用于写入 Word
            original_text: 原始完整文本，用于审计对比
        num_id_map: {para_idx: num_id} 映射（如果为空则自动创建）
        format_type: 格式类型
    """
    # 收集 heading 样式名称
    heading_styles = {'Heading ' + str(i) for i in range(1, 10)}

    # 构建已编号段落的索引映射
    numbered_map = {idx: (level, clean_text, original_text)
                    for idx, level, clean_text, original_text in numbered_paras}
    numbered_indices = set(numbered_map.keys())

    # 提取所有 Level 0 索引并排序
    level0_indices = sorted(idx for idx, level, _, _ in numbered_paras if level == 0)

    # ============================================================
    # 阶段1：对一级标题之间的段落，自动推断级别并补充
    # ============================================================
    extra_paras = {}  # idx -> (level, clean_text, original_text)

    from .detector import is_signature_section

    for i in level0_indices:
        next_idx = level0_indices[level0_indices.index(i) + 1] if level0_indices.index(i) < len(level0_indices) - 1 else len(doc.paragraphs)

        in_signature_zone = False

        for j in range(i + 1, next_idx):
            if j in numbered_indices:
                continue

            para = doc.paragraphs[j]
            text = para.text.strip()

            if not text:
                continue

            # 签名区检测
            if is_signature_section(text):
                in_signature_zone = True
                continue

            if in_signature_zone:
                continue

            # 纯标题行（如"甲方信息"）-> 不编号
            if _should_skip_paragraph(text):
                continue

            # 推断为二级（level 1）
            extra_paras[j] = (1, text, text)

    # 合并额外推断的段落到 numbered_paras
    if extra_paras:
        all_paras = []
        all_idx_set = set(numbered_indices) | set(extra_paras.keys())
        for j in sorted(all_idx_set):
            if j in numbered_map:
                all_paras.append((j, numbered_map[j][0], numbered_map[j][1], numbered_map[j][2]))
            else:
                all_paras.append((j, extra_paras[j][0], extra_paras[j][1], extra_paras[j][2]))
        numbered_paras = all_paras
        numbered_map = {idx: (level, clean, orig)
                        for idx, level, clean, orig in numbered_paras}

    # ============================================================
    # 阶段2：清除所有不被编号段落的旧 numPr
    # ============================================================
    all_numbered_indices = {idx for idx, _, _, _ in numbered_paras}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if i in all_numbered_indices:
            continue
        _remove_numPr(para)

    # ============================================================
    # 阶段3：创建新的编号定义并应用
    # ============================================================
    numbering_part, numbering_elem = create_numbering_part(doc)
    abstract_id = next_abstract_id(numbering_elem)
    create_abstract_numbering(numbering_elem, abstract_id=abstract_id, format_type=format_type)

    level0_indices = [idx for idx, level, _, _ in numbered_paras if level == 0]
    num_id_map = create_num_instances(numbering_elem, abstract_id=abstract_id,
                                      level0_indices=level0_indices, format_type=format_type)

    # 更新 numbering part
    if hasattr(numbering_part, '_blob'):
        numbering_part._blob = etree.tostring(numbering_elem, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 应用自动编号
    current_num_id = None

    for para_idx, level, clean_text, original_text in numbered_paras:
        para = doc.paragraphs[para_idx]

        # 确定 numId
        if level == 0:
            current_num_id = num_id_map.get(para_idx)

        # 应用自动编号
        num_id_to_use = current_num_id if level >= 1 else num_id_map.get(para_idx)
        apply_numbering_to_paragraph(para, level, num_id_to_use, clean_text)


def verify_numbering(doc: Document, numbered_paras: list[tuple[int, int, str, str]], *, original_doc: Document | None = None, format_type: str = 'chinese') -> dict:
    """验证自动编号是否正确应用（增强版，含审计检查）

    Args:
        doc: 输出文档
        numbered_paras: 编号段落列表
        original_doc: 原始文档（用于审计对比）
        format_type: 编号格式类型

    Returns:
        验证结果字典
    """
    from .auditor import audit_completeness
    from .detector import is_signature_section

    results = []
    all_valid = True

    # 创建已编号段落的索引集合
    numbered_indices = {idx for idx, _, _, _ in numbered_paras}

    # 验证所有段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        p_pr = para._element.find(qn('w:pPr'))
        has_num = False
        actual_level = None
        num_id = None

        if p_pr is not None:
            num_pr = p_pr.find(qn('w:numPr'))
            if num_pr is not None:
                has_num = True
                ilvl_elem = num_pr.find(qn('w:ilvl'))
                num_id_elem = num_pr.find(qn('w:numId'))
                if ilvl_elem is not None:
                    actual_level = int(ilvl_elem.get(qn('w:val'), -1))
                if num_id_elem is not None:
                    num_id = num_id_elem.get(qn('w:val'))

        # 检查是否是签名部分
        is_signature = is_signature_section(text)

        # 判断是否有效
        if i in numbered_indices:
            # 应该编号的段落
            expected_level = next(level for idx, level, _, _ in numbered_paras if idx == i)
            is_valid = has_num and actual_level == expected_level
        elif is_signature:
            # 签名部分不应该编号
            is_valid = not has_num
        else:
            # 其他段落：允许无编号（如普通正文）
            is_valid = True

        results.append({
            'para_idx': i,
            'expected_level': expected_level if i in numbered_indices else None,
            'actual_level': actual_level,
            'has_num': has_num,
            'num_id': num_id,
            'is_signature': is_signature,
            'text': text[:50],
            'valid': is_valid,
        })

        if not is_valid:
            all_valid = False

    # 运行增强审计（优先使用，更强的可靠性）
    audit_report = None
    if original_doc is not None:
        audit_report = audit_completeness(original_doc, doc, format_type, numbered_paras=numbered_paras)
        # 审计报告优先：即使 basic verification 说 "all_valid"，
        # 只要审计发现遗漏或层级问题，就覆盖为 not valid
        if not audit_report.all_clear:
            all_valid = False
            logger.warning("审计发现问题，recommend review:\n%s", audit_report.summary)

    return {
        'all_valid': all_valid,
        'total': len(results),
        'valid_count': sum(1 for r in results if r['valid']),
        'invalid_count': sum(1 for r in results if not r['valid']),
        'results': [r for r in results if not r['valid']],  # 只返回失败的
        'audit': audit_report._asdict() if audit_report else None,
    }
