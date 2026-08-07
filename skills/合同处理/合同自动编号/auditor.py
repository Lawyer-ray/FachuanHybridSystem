"""
审计模块

提供转换后自动审计检查，利用 AI 大模型能力发现遗漏和错误。

设计原则：
1. 对比原始文档和输出文档，找出所有差异
2. 检测编号序列的连续性（如一、二、三 → 是否缺四）
3. 检测编号前缀 vs 自动编号的匹配关系
4. 生成结构化审计报告，供 AI 进一步检查和提醒
"""

from __future__ import annotations

import logging
import re
from typing import NamedTuple

from docx import Document

from .detector import detect_chinese_level0, detect_chinese_sublevel, detect_decimal_level0, detect_decimal_sublevel

logger = logging.getLogger(__name__)

# ============================================================================
# 数据结构
# ============================================================================


class ParagraphAudit(NamedTuple):
    """单段审计结果"""
    para_idx: int
    original_text: str
    output_text: str
    has_numid: bool
    expected_number_prefix: str | None  # 检测到的手动编号前缀
    issue: str | None  # 如果有问题，描述问题


class AuditReport(NamedTuple):
    """完整审计报告"""
    total_paragraphs: int  # 总段落数
    numbered_paras: int  # 预期编号段落数
    missed_paras: list[ParagraphAudit]  # 遗漏的段落
    mismatch_paras: list[ParagraphAudit]  # 编号不匹配段落
    sequence_gaps: list[dict]  # 编号序列缺口
    hierarchy_issues: list[dict]  # 层级结构问题
    all_clear: bool  # 是否全部通过
    summary: str  # AI 友好的摘要文本


# ============================================================================
# 核心检测
# ============================================================================


def _extract_number_prefix(text: str, format_type: str) -> str | None:
    """从文本中提取手动编号前缀

    例如：
      "一、合作原则" → "一、"
      "（四）本协议自..." → "（四）"
      "1.配合工作" → "1."
    """
    # 中文格式一级
    m = detect_chinese_level0(text)
    if m[0]:
        return m[0]

    # 纯数字格式一级
    m = detect_decimal_level0(text)
    if m[0]:
        return m[0]

    # 子级编号（中文）
    m = detect_chinese_sublevel(text, has_level1_heading=True)
    if m[0] is not None:
        return m[1]

    # 子级编号（纯数字）
    m = detect_decimal_sublevel(text)
    if m[0] is not None:
        return m[1]

    return None


def _has_number_prefix(text: str) -> bool:
    """检查文本行首是否有编号前缀（无论格式类型）"""
    return _extract_number_prefix(text, 'chinese') is not None


def _detect_sequence_gaps(original_doc: Document) -> list[dict]:
    """检测原始文档中编号序列的潜在缺口

    例如文档有一、二、三，但缺四 → 提醒检查

    Returns:
        每个缺口的信息
    """
    gaps = []

    # 收集所有一级标题的中文序号（一、二、三...）
    chinese_nums = []
    for i, para in enumerate(original_doc.paragraphs):
        text = para.text.strip()
        m = re.match(r'^([一二三四五六七八九十]+)、', text)
        if m:
            chinese_nums.append({
                'para_idx': i,
                'char': m.group(1),
                'text': text,
            })

    # 极简检查：如果有一、二、三，但没有四，且文本后面有"五"
    # → 很可能是缺口（也可能是文档结构变异，作为提示而非错误）
    # 更可靠的检查：遍历收集到的序号，看是否完整序列

    # 收集所有序号值（从数字到序号的映射）
    char_to_num = {
        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
        '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
    }

    if len(chinese_nums) >= 2:
        nums = []
        for item in chinese_nums:
            cn = item['char']
            # 简单处理：目前只支持一到十
            n = char_to_num.get(cn)
            if n:
                nums.append({**item, 'num': n})

        # 检测连续序列中的缺失
        for i in range(1, len(nums)):
            prev = nums[i - 1]['num']
            curr = nums[i]['num']
            if curr > prev + 1:
                # 有序号缺失
                missing = [char_to_num.get(k) for k in char_to_num if prev < (char_to_num.get(k) or 0) < curr]
                if missing:
                    gaps.append({
                        'before_para': nums[i - 1]['para_idx'],
                        'before_char': nums[i - 1]['char'],
                        'after_char': nums[i]['char'],
                        'missing': [list(char_to_num.keys())[list(char_to_num.values()).index(m)] for m in missing],
                        'hint': f"一级标题 {nums[i-1]['char']}、与 {nums[i]['char']}、之间缺少序号：{''.join([list(char_to_num.keys())[list(char_to_num.values()).index(m)] + '、' for m in missing])}请检查是否有遗漏章节",
                    })

    return gaps


def validate_hierarchy(numbered_paras: list[tuple[int, int, str, str]]) -> list[dict]:
    """验证编号层级结构的合理性

    检查规则：
    - （一）型子标题（level 1，matched 以（开头）后的无编号段落应在 level 2
    - 如果子标题后的段落仍在 level 1，标记为层级错误

    Args:
        numbered_paras: 编号段落列表 (para_idx, level, matched, text)

    Returns:
        层级问题列表，每项含 para_idx, heading_idx, issue, text
    """
    issues = []

    for i, (para_idx, level, matched, text) in enumerate(numbered_paras):
        # 检测（一）型子标题（level 1，matched 以（或(开头）
        is_subheading = level == 1 and matched and matched[0] in '（('

        if not is_subheading:
            continue

        # 检查后续段落，直到遇到下一个同级或更高级标题
        for j in range(i + 1, len(numbered_paras)):
            next_idx, next_level, next_matched, next_text = numbered_paras[j]

            # 遇到更高级别标题（level 更小），停止检查
            if next_level < level:
                break

            # 遇到同级子标题（有编号前缀），停止检查
            if next_level == level and next_matched:
                break

            # 同级无编号内容段落 → 层级错误（应在 level+1）
            if next_level == level and not next_matched:
                issues.append({
                    'para_idx': next_idx,
                    'heading_idx': para_idx,
                    'heading_text': text[:50],
                    'issue': f'段落在子标题"{text[:30]}"之后，'
                             f'但层级相同（L{level}），应为 L{level + 1}',
                    'text': next_text[:50],
                })

    return issues


def audit_completeness(
    original_doc: Document,
    output_doc: Document,
    format_type: str,
    numbered_paras: list[tuple[int, int, str, str]] | None = None,
) -> AuditReport:
    """审计输出文档的编号完整性

    核心检查项：
    1. 原文件中所有带编号前缀的段落，输出文件中是否都有自动编号
    2. 输出中有自动编号的段落，文本是否与原始文档一致（编号前缀被替换为自动编号）
    3. 编号序列是否连续
    4. 层级结构是否合理（子标题后的内容应在更深层级）

    Args:
        original_doc: 原始文档
        output_doc: 输出文档
        format_type: 编号格式类型
        numbered_paras: 编号段落列表（可选，用于层级结构验证）

    Returns:
        审计报告
    """
    # 分析原文件：所有带编号前缀的段落
    original_numbered: dict[int, str] = {}  # para_idx → text
    for i, para in enumerate(original_doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        prefix = _extract_number_prefix(text, format_type)
        if prefix:
            original_numbered[i] = text

    # 分析输出文件：所有带 numId 的段落
    output_with_num: dict[int, str] = {}
    for i, para in enumerate(output_doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # 检查是否有 numId (自动编号)
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        num_pr = para._element.find(f'{{{ns}}}pPr/{{{ns}}}numPr')
        if num_pr is not None:
            output_with_num[i] = text

    # ------------------------------------------------------------------
    # 检查1: 原文件有编号的段落，输出是否也有编号
    # ------------------------------------------------------------------
    missed: list[ParagraphAudit] = []
    for i, orig_text in original_numbered.items():
        if i >= len(output_doc.paragraphs):
            missed.append(ParagraphAudit(
                para_idx=i,
                original_text=orig_text,
                output_text='(段落不存在)',
                has_numid=False,
                expected_number_prefix=_extract_number_prefix(orig_text, format_type),
                issue='原文件带编号的段落，在输出文件中不存在',
            ))
            continue

        out_para = output_doc.paragraphs[i]
        out_text = out_para.text.strip()
        has_num = _has_numid(out_para)

        if not has_num:
            issue = (
                f'原文件段落 [{i}] 有编号前缀 "{_extract_number_prefix(orig_text, format_type)}"，'
                f'但输出文件未设置自动编号。文本: {orig_text[:60]}'
            )
            missed.append(ParagraphAudit(
                para_idx=i,
                original_text=orig_text,
                output_text=out_text,
                has_numid=False,
                expected_number_prefix=_extract_number_prefix(orig_text, format_type),
                issue=issue,
            ))

    # ------------------------------------------------------------------
    # 检查2: 输出文件有编号的段落，文本中不应该残留手动编号前缀
    # ------------------------------------------------------------------
    mismatched: list[ParagraphAudit] = []
    for i, out_text in output_with_num.items():
        # 检查输出文本中是否还有手动编号前缀
        prefix = _extract_number_prefix(out_text, format_type)
        if prefix:
            issue = (
                f'输出文件段落 [{i}] 同时有自动编号和手动编号残留 "{prefix}"，'
                f'可能导致双重编号显示。文本: {out_text[:60]}'
            )
            mismatched.append(ParagraphAudit(
                para_idx=i,
                original_text=original_numbered.get(i, '(原文件无此编号段落，可能是新增)'),
                output_text=out_text,
                has_numid=True,
                expected_number_prefix=prefix,
                issue=issue,
            ))

    # ------------------------------------------------------------------
    # 检查3: 编号序列连续性
    # ------------------------------------------------------------------
    sequence_gaps = _detect_sequence_gaps(original_doc)

    # ------------------------------------------------------------------
    # 检查4: 层级结构合理性
    # ------------------------------------------------------------------
    hierarchy_issues: list[dict] = []
    if numbered_paras is not None:
        hierarchy_issues = validate_hierarchy(numbered_paras)

    # ------------------------------------------------------------------
    # 生成摘要
    # ------------------------------------------------------------------
    all_clear = (
        len(missed) == 0
        and len(mismatched) == 0
        and len(hierarchy_issues) == 0
    )

    lines = []
    lines.append(f'审计完成：{len(original_numbered)} 个原文件编号段落，'
                 f'{len(output_with_num)} 个输出自动编号段落。')

    if all_clear:
        lines.append('✿ 全部通过，未发现遗漏或残留。')
    else:
        if missed:
            lines.append(f'⚠ 遗漏 {len(missed)} 个编号段落（原文件有编号但输出未设置自动编号）：')
            for m in missed:
                lines.append(f'   [{m.para_idx}] 前缀={m.expected_number_prefix} → {m.original_text[:50]}')

        if mismatched:
            lines.append(f'⚠ {len(mismatched)} 个段落残留手动编号前缀：')
            for m in mismatched:
                lines.append(f'   [{m.para_idx}] {m.output_text[:50]}')

        if hierarchy_issues:
            lines.append(f'⚠ 层级结构 {len(hierarchy_issues)} 处问题（子标题后内容层级错误）：')
            for h in hierarchy_issues:
                lines.append(f'   [{h["para_idx"]}] {h["issue"]}: {h["text"]}')

    if sequence_gaps:
        lines.append(f'ℹ 编号序列 {len(sequence_gaps)} 处可能不连续：')
        for g in sequence_gaps:
            lines.append(f'   {g["hint"]}')

    summary = '\n'.join(lines)

    return AuditReport(
        total_paragraphs=len(output_doc.paragraphs),
        numbered_paras=len(original_numbered),
        missed_paras=missed,
        mismatch_paras=mismatched,
        sequence_gaps=sequence_gaps,
        hierarchy_issues=hierarchy_issues,
        all_clear=all_clear,
        summary=summary,
    )


def _has_numid(para) -> bool:
    """检查段落是否有 numId（自动编号标记）"""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    num_pr = para._element.find(f'{{{ns}}}pPr/{{{ns}}}numPr')
    return num_pr is not None and num_pr.find(f'{{{ns}}}numId') is not None
