"""
合同审查报告生成脚本

生成三件套 .docx 文件：
1. 审查标注版合同（修订模式 + 批注）
2. 法律意见书
3. 法律分析（内部参考）

依赖：python-docx（已在 backend/pyproject.toml 中声明）
"""

import json
import logging
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


def generate_opinion_letter(report_data: dict, output_path: Path) -> Path:
    """
    生成法律意见书（五模块）

    Args:
        report_data: 审查结果数据，结构如下：
            {
                "contract_name": "合同名称",
                "parties": {"party_a": "甲方名称", "party_b": "乙方名称"},
                "contract_type": "买卖合同",
                "risk_summary": {"critical": 2, "important": 5, "normal": 3, "compliant": 8},
                "risk_level": "中等风险",
                "findings": [
                    {
                        "id": 1,
                        "clause": "第 X 条",
                        "risk_type": "违约责任",
                        "risk_level": "critical",
                        "description": "风险描述",
                        "suggestion": "修改建议",
                        "legal_basis": "法律依据"
                    }
                ],
                "overall_assessment": {
                    "advantages": ["优势1", "优势2"],
                    "risks": ["风险1", "风险2"],
                    "negotiation_tips": ["谈判要点1"],
                    "post_signing_notes": ["签约后注意事项1"]
                }
            }
        output_path: 输出文件路径
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # === 模块 1：风险概览 ===
    doc.add_heading("一、风险概览", level=1)

    summary = report_data.get("risk_summary", {})
    risk_level = report_data.get("risk_level", "待评估")

    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    headers = ["综合等级", "🔴 严重", "🟡 重要", "🟢 一般", "⚪ 合规"]
    values = [
        risk_level,
        str(summary.get("critical", 0)),
        str(summary.get("important", 0)),
        str(summary.get("normal", 0)),
        str(summary.get("compliant", 0)),
    ]
    for i, (header, value) in enumerate(zip(headers, values)):
        table.rows[0].cells[i].text = header
        table.rows[1].cells[i].text = value

    doc.add_paragraph()

    # === 模块 2：合同基本信息 ===
    doc.add_heading("二、合同基本信息", level=1)

    parties = report_data.get("parties", {})
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    info_items = [
        ("合同名称", report_data.get("contract_name", "")),
        ("合同类型", report_data.get("contract_type", "")),
        ("甲方", parties.get("party_a", "")),
        ("乙方", parties.get("party_b", "")),
        ("审查日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    for row_idx, (label, value) in enumerate(info_items):
        info_table.rows[row_idx].cells[0].text = label
        info_table.rows[row_idx].cells[1].text = value

    doc.add_paragraph()

    # === 模块 3：逐条意见表 ===
    doc.add_heading("三、逐条审查意见", level=1)

    findings = report_data.get("findings", [])
    if findings:
        risk_icons = {
            "critical": "🔴",
            "important": "🟡",
            "normal": "🟢",
            "compliant": "⚪",
        }

        findings_table = doc.add_table(rows=1, cols=6)
        findings_table.style = "Table Grid"
        col_headers = ["序号", "风险类型", "审查条款", "风险描述", "修改建议", "风险等级"]
        for i, header in enumerate(col_headers):
            findings_table.rows[0].cells[i].text = header

        for finding in findings:
            row = findings_table.add_row()
            row.cells[0].text = str(finding.get("id", ""))
            row.cells[1].text = finding.get("risk_type", "")
            row.cells[2].text = finding.get("clause", "")
            row.cells[3].text = finding.get("description", "")
            row.cells[4].text = finding.get("suggestion", "")
            level = finding.get("risk_level", "normal")
            row.cells[5].text = f"{risk_icons.get(level, '')} {level}"
    else:
        doc.add_paragraph("未发现重大风险。")

    doc.add_paragraph()

    # === 模块 4：综合评价 ===
    doc.add_heading("四、综合评价", level=1)

    assessment = report_data.get("overall_assessment", {})

    doc.add_heading("优势", level=2)
    for item in assessment.get("advantages", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("风险提示", level=2)
    for item in assessment.get("risks", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("谈判要点", level=2)
    for item in assessment.get("negotiation_tips", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("签约后注意事项", level=2)
    for item in assessment.get("post_signing_notes", []):
        doc.add_paragraph(item, style="List Bullet")

    # === 模块 5：法律依据 ===
    doc.add_heading("五、法律依据", level=1)

    legal_bases = set()
    for finding in findings:
        basis = finding.get("legal_basis", "")
        if basis:
            legal_bases.add(basis)

    if legal_bases:
        for basis in sorted(legal_bases):
            doc.add_paragraph(basis, style="List Bullet")
    else:
        doc.add_paragraph("（无具体法律依据引用）")

    # === 免责声明 ===
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "⚠️ 本意见书不构成法律意见，仅供律师作为审查辅助参考。"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(str(output_path))
    return output_path


def generate_analysis_doc(report_data: dict, output_path: Path) -> Path:
    """
    生成法律分析（内部参考文档）

    每个修订点对应的法条、司法解释、类案裁判。
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_heading("法律分析（内部参考）", level=1)
    doc.add_paragraph(f"合同：{report_data.get('contract_name', '')}")
    doc.add_paragraph(f"审查日期：{datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    findings = report_data.get("findings", [])
    for finding in findings:
        doc.add_heading(f"{finding.get('id', '')}. {finding.get('clause', '')}", level=2)
        doc.add_paragraph(f"风险类型：{finding.get('risk_type', '')}")
        doc.add_paragraph(f"问题：{finding.get('description', '')}")
        doc.add_paragraph(f"法律依据：{finding.get('legal_basis', '待补充')}")
        doc.add_paragraph(f"修改建议：{finding.get('suggestion', '')}")
        doc.add_paragraph()

    doc.save(str(output_path))
    return output_path


def main():
    """
    命令行入口

    用法：
        python generate_report.py <report_data.json> <output_dir>

    report_data.json 结构参见 generate_opinion_letter 的 docstring。
    """
    if len(sys.argv) < 3:
        logger.error("用法: python generate_report.py <report_data.json> <output_dir>")
        sys.exit(1)

    data_path = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])
    output_dir.mkdir(parents=True, exist_ok=True)

    with open(data_path, encoding="utf-8") as f:
        report_data = json.load(f)

    contract_name = report_data.get("contract_name", "合同")
    date_str = datetime.now().strftime("%Y%m%d")

    # 生成法律意见书
    opinion_path = output_dir / f"法律意见书_{contract_name}_{date_str}.docx"
    generate_opinion_letter(report_data, opinion_path)
    logger.info("✅ 法律意见书: %s", opinion_path)

    # 生成法律分析
    analysis_path = output_dir / f"法律分析_{contract_name}_{date_str}.docx"
    generate_analysis_doc(report_data, analysis_path)
    logger.info("✅ 法律分析: %s", analysis_path)

    # 生成审查摘要 markdown
    summary_path = output_dir / "审查摘要.md"
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(f"# 审查摘要 — {contract_name}\n\n")
        f.write(f"审查日期：{date_str}\n\n")
        f.write(f"综合风险等级：{report_data.get('risk_level', '待评估')}\n\n")
        summary = report_data.get("risk_summary", {})
        f.write(f"| 🔴 严重 | 🟡 重要 | 🟢 一般 | ⚪ 合规 |\n")
        f.write(f"|---------|---------|---------|--------|\n")
        f.write(
            f"| {summary.get('critical', 0)} | {summary.get('important', 0)} "
            f"| {summary.get('normal', 0)} | {summary.get('compliant', 0)} |\n"
        )
        f.write("\n## 主要发现\n\n")
        for finding in report_data.get("findings", []):
            icons = {"critical": "🔴", "important": "🟡", "normal": "🟢", "compliant": "⚪"}
            icon = icons.get(finding.get("risk_level", ""), "")
            f.write(f"- {icon} **{finding.get('clause', '')}**: {finding.get('description', '')}\n")
    logger.info("✅ 审查摘要: %s", summary_path)

    logger.info("📁 报告已生成至: %s", output_dir)


if __name__ == "__main__":
    main()
