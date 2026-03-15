#!/usr/bin/env python3
"""创建模拟庭审报告Word模板."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_template():
    """创建模拟庭审报告模板."""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('模拟庭审报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表格
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    headers = [
        ('案件名称', '{{模拟庭审_案件名称}}'),
        ('案由', '{{模拟庭审_案由}}'),
        ('模拟模式', '{{模拟庭审_模式}}'),
        ('生成时间', '{{模拟庭审_生成时间}}'),
    ]

    for i, (label, value) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    # 争议焦点
    doc.add_paragraph()
    doc.add_heading('一、争议焦点', level=1)
    p = doc.add_paragraph()
    p.add_run('{{模拟庭审_争议焦点}}')

    # 证据分析
    doc.add_paragraph()
    doc.add_heading('二、证据分析', level=1)
    p = doc.add_paragraph()
    p.add_run('{{模拟庭审_证据分析}}')

    # 风险评估
    doc.add_paragraph()
    doc.add_heading('三、风险评估', level=1)
    p = doc.add_paragraph()
    p.add_run('{{模拟庭审_风险评估}}')

    # 胜诉概率
    doc.add_paragraph()
    doc.add_heading('四、胜诉概率', level=1)
    p = doc.add_paragraph()
    p.add_run('{{模拟庭审_胜诉概率}}')

    # 建议策略
    doc.add_paragraph()
    doc.add_heading('五、建议策略', level=1)
    p = doc.add_paragraph()
    p.add_run('{{模拟庭审_建议策略}}')

    # 保存
    output_path = '/Users/huangsong21/Downloads/Coding/AI/FachuanHybridSystem/backend/apps/documents/docx_templates/2-案件材料/5-模拟庭审报告/模拟庭审报告.docx'
    doc.save(output_path)
    print(f'模板创建成功: {output_path}')


if __name__ == '__main__':
    create_template()
