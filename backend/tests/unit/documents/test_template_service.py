# -*- coding: utf-8 -*-
"""
DocumentTemplateService 单元测试

测试文书模板服务的核心功能：
- 模板创建（上传模式和路径引用模式）
- 文件路径验证
- 版本管理
- 占位符提取
"""

import os
import tempfile
from unittest.mock import MagicMock, patch

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile

from apps.core.exceptions import NotFoundError, ValidationException
from apps.documents.models import DocumentTemplate, DocumentTemplateType, Placeholder, PlaceholderCategory
from apps.documents.services.template_service import DocumentTemplateService


@pytest.fixture
def template_service():
    """创建 DocumentTemplateService 实例"""
    return DocumentTemplateService()


@pytest.fixture
def temp_docx_file():
    """创建临时 docx 文件用于测试"""
    # 使用 docxtpl 创建模板文件
    from docx import Document
    from docxtpl import DocxTemplate

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        # 先创建基础文档
        doc = Document()
        doc.add_paragraph("测试文档 {{ case_name }}")
        doc.add_paragraph("当事人: {{ party_name }}")
        doc.save(f.name)
        temp_path = f.name

    yield temp_path

    # 清理
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def uploaded_file(temp_docx_file):
    """创建上传文件对象"""
    with open(temp_docx_file, "rb") as f:
        content = f.read()
    return SimpleUploadedFile(
        name="test_template.docx",
        content=content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


class TestDocumentTemplateServiceCreate:
    """测试模板创建功能"""

    @pytest.mark.django_db
    def test_create_template_with_file_path(self, template_service, temp_docx_file):
        """测试使用文件路径创建模板"""
        template = template_service.create_template(
            name="测试模板",
            template_type=DocumentTemplateType.CASE,
            file_path=temp_docx_file,
            description="测试描述",
            case_types=["civil", "criminal"],
        )

        assert template.id is not None
        assert template.name == "测试模板"
        assert template.template_type == DocumentTemplateType.CASE
        assert template.file_path == temp_docx_file
        assert template.is_active is True

    @pytest.mark.django_db
    def test_create_template_with_uploaded_file(self, template_service, uploaded_file):
        """测试使用上传文件创建模板"""
        template = template_service.create_template(
            name="上传模板", template_type=DocumentTemplateType.CONTRACT, file=uploaded_file
        )

        assert template.id is not None
        assert template.name == "上传模板"
        assert template.file is not None
        assert template.file_path == ""

    @pytest.mark.django_db
    def test_create_template_without_file_raises_error(self, template_service):
        """测试不提供文件时抛出异常"""
        with pytest.raises(ValidationException) as exc_info:
            template_service.create_template(name="无文件模板", template_type=DocumentTemplateType.CONTRACT)

        assert "必须提供上传文件或文件路径" in str(exc_info.value)

    @pytest.mark.django_db
    def test_create_template_with_both_file_and_path_raises_error(
        self, template_service, uploaded_file, temp_docx_file
    ):
        """测试同时提供文件和路径时抛出异常"""
        with pytest.raises(ValidationException) as exc_info:
            template_service.create_template(
                name="双重模板",
                template_type=DocumentTemplateType.CONTRACT,
                file=uploaded_file,
                file_path=temp_docx_file,
            )

        assert "不能同时提供" in str(exc_info.value)

    @pytest.mark.django_db
    def test_create_template_with_invalid_path_raises_error(self, template_service):
        """测试使用不存在的路径时抛出异常"""
        with pytest.raises(ValidationException) as exc_info:
            template_service.create_template(
                name="无效路径模板",
                template_type=DocumentTemplateType.CONTRACT,
                file_path="/nonexistent/path/template.docx",
            )

        assert "文件不存在" in str(exc_info.value)


class TestDocumentTemplateServiceValidation:
    """测试文件路径验证功能"""

    def test_validate_file_path_with_existing_file(self, template_service, temp_docx_file):
        """测试验证存在的文件路径"""
        assert template_service.validate_file_path(temp_docx_file) is True

    def test_validate_file_path_with_nonexistent_file(self, template_service):
        """测试验证不存在的文件路径"""
        assert template_service.validate_file_path("/nonexistent/file.docx") is False

    def test_validate_file_path_with_empty_path(self, template_service):
        """测试验证空路径"""
        assert template_service.validate_file_path("") is False
        assert template_service.validate_file_path(None) is False


class TestDocumentTemplateServiceVersioning:
    """测试版本管理功能"""

    @pytest.mark.django_db
    def test_update_template_file(self, template_service, temp_docx_file):
        """测试更新模板文件"""
        # 创建初始模板
        template = template_service.create_template(
            name="版本测试模板", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file
        )

        # 创建新的临时文件
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("更新后的文档")
            doc.save(f.name)
            new_file_path = f.name

        try:
            # 更新模板文件
            updated = template_service.update_template(template_id=template.id, file_path=new_file_path)

            assert updated.file_path == new_file_path
        finally:
            if os.path.exists(new_file_path):
                os.unlink(new_file_path)

    @pytest.mark.django_db
    def test_update_template_name_only(self, template_service, temp_docx_file):
        """测试只更新名称字段"""
        template = template_service.create_template(
            name="版本测试模板2", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file
        )

        # 只更新名称
        updated = template_service.update_template(template_id=template.id, name="新名称")

        assert updated.name == "新名称"

    @pytest.mark.django_db
    def test_update_nonexistent_template_raises_error(self, template_service):
        """测试更新不存在的模板时抛出异常"""
        with pytest.raises(NotFoundError):
            template_service.update_template(template_id=99999, name="新名称")


class TestDocumentTemplateServicePlaceholders:
    """测试占位符提取功能"""

    @pytest.mark.django_db
    def test_extract_placeholders_from_docx(self, template_service, temp_docx_file):
        """测试从 docx 文件提取占位符"""
        template = template_service.create_template(
            name="占位符测试模板", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file
        )

        placeholders = template_service.extract_placeholders(template)

        # 应该提取到 case_name 和 party_name
        assert "case_name" in placeholders
        assert "party_name" in placeholders

    @pytest.mark.django_db
    def test_get_undefined_placeholders(self, template_service, temp_docx_file):
        """测试获取未定义的占位符"""
        # 创建模板
        template = template_service.create_template(
            name="未定义占位符测试", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file
        )

        # 注册一个占位符
        Placeholder.objects.create(  # type: ignore[misc]
            key="case_name", display_name="案件名称", data_path="case.name", category=PlaceholderCategory.CASE
        )

        undefined = template_service.get_undefined_placeholders(template)

        # party_name 未注册，应该在未定义列表中
        assert "party_name" in undefined
        assert "case_name" not in undefined

    @pytest.mark.django_db
    def test_extract_placeholders_from_nonexistent_file(self, template_service):
        """测试从不存在的文件提取占位符返回空列表"""
        # 创建一个模板但文件路径无效
        template = DocumentTemplate.objects.create(
            name="无效文件模板", template_type=DocumentTemplateType.CONTRACT, file_path="/nonexistent/file.docx"
        )

        placeholders = template_service.extract_placeholders(template)
        assert placeholders == []


class TestDocumentTemplateServiceQuery:
    """测试查询功能"""

    @pytest.mark.django_db
    def test_get_template_by_id(self, template_service, temp_docx_file):
        """测试根据 ID 获取模板"""
        template = template_service.create_template(
            name="查询测试模板", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file
        )

        found = template_service.get_template_by_id(template.id)
        assert found.id == template.id
        assert found.name == "查询测试模板"

    @pytest.mark.django_db
    def test_get_template_by_id_not_found(self, template_service):
        """测试获取不存在的模板"""
        with pytest.raises(NotFoundError):
            template_service.get_template_by_id(99999)

    @pytest.mark.django_db
    def test_list_templates_with_filters(self, template_service, temp_docx_file):
        """测试带过滤条件的列表查询"""
        # 创建多个模板
        template_service.create_template(
            name="案件模板1", template_type=DocumentTemplateType.CASE, file_path=temp_docx_file, case_types=["civil"]
        )
        template_service.create_template(
            name="合同模板1", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file, is_active=False
        )

        # 按类型过滤
        case_templates = template_service.list_templates(template_type=DocumentTemplateType.CASE)
        assert len(case_templates) >= 1
        assert all(t.template_type == DocumentTemplateType.CASE for t in case_templates)

        # 按启用状态过滤
        active_templates = template_service.list_templates(is_active=True)
        assert all(t.is_active for t in active_templates)

    @pytest.mark.django_db
    def test_delete_template_soft_delete(self, template_service, temp_docx_file):
        """测试软删除模板"""
        template = template_service.create_template(
            name="删除测试模板", template_type=DocumentTemplateType.CONTRACT, file_path=temp_docx_file
        )

        result = template_service.delete_template(template.id)
        assert result is True

        # 模板仍然存在但已禁用
        template.refresh_from_db()
        assert template.is_active is False
