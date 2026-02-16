"""
补充协议生成功能集成测试

测试完整的补充协议生成流程，包括：
- 正常生成流程
- 错误处理场景
- 模板查找逻辑
- 文件名生成
- 占位符替换

Requirements: 2.3
"""

import os
import tempfile
from decimal import Decimal

import pytest
from django.test import TestCase
from django.utils import timezone
from docx import Document

from apps.contracts.models import (
    Contract,
    ContractParty,
    PartyRole,
    SupplementaryAgreement,
    SupplementaryAgreementParty,
)
from apps.core.enums import CaseType
from apps.documents.models import DocumentContractSubType, DocumentTemplateType
from apps.documents.services.generation.supplementary_agreement_generation_service import (
    SupplementaryAgreementGenerationService,
)
from tests.factories.client_factories import ClientFactory
from tests.factories.contract_factories import ContractAssignmentFactory, ContractFactory
from tests.factories.document_factories import DocumentTemplateFactory
from tests.factories.organization_factories import LawFirmFactory, LawyerFactory


@pytest.mark.django_db
class TestSupplementaryAgreementGenerationIntegration:
    """补充协议生成功能集成测试"""

    def setup_method(self):
        """每个测试方法前执行"""
        self.service = SupplementaryAgreementGenerationService()

        # 创建测试用户和律所
        self.law_firm = LawFirmFactory()
        self.lawyer = LawyerFactory(law_firm=self.law_firm)

        # 创建测试客户
        self.client1 = ClientFactory(name="张三", client_type="individual", id_number="110101199001011234")
        self.client2 = ClientFactory(name="李四", client_type="individual", id_number="110101199002022345")
        self.client3 = ClientFactory(name="王五公司", client_type="company", id_number="91110000123456789X")

        # 创建测试合同
        self.contract = ContractFactory(
            name="张三与李四一案代理合同", case_type=CaseType.CIVIL, specified_date=timezone.localdate()
        )

        # 创建合同律师指派
        ContractAssignmentFactory(contract=self.contract, lawyer=self.lawyer, is_primary=True)

        # 创建合同当事人（原合同只有张三）
        ContractParty.objects.create(contract=self.contract, client=self.client1, role=PartyRole.PRINCIPAL)

    def _create_test_template(self, case_type="civil"):
        """创建测试模板文件"""
        # 创建一个简单的DOCX文件用于测试
        from docx import Document

        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_file.close()

        # 创建DOCX文档
        doc = Document()
        doc.add_paragraph("补充协议")
        doc.add_paragraph("{{补充协议名称}}")
        doc.add_paragraph("甲方：{{委托人信息}}")
        doc.add_paragraph("{{委托人主体信息条款}}")
        doc.add_paragraph("{{对方当事人主体信息条款}}")
        doc.add_paragraph("签名盖章：")
        doc.add_paragraph("{{委托人签名盖章信息}}")
        doc.add_paragraph("{{年份}}年")

        # 保存文档
        doc.save(temp_file.name)

        # 创建文档模板记录
        template = DocumentTemplateFactory(
            name="补充协议模板",
            template_type=DocumentTemplateType.CONTRACT,
            contract_sub_type=DocumentContractSubType.SUPPLEMENTARY_AGREEMENT,
            contract_types=[case_type],
            is_active=True,
            file_path=temp_file.name,
        )

        return template, temp_file.name

    def _create_supplementary_agreement(self, name="补充协议一", principals=None, opposing=None):
        """创建补充协议"""
        agreement = SupplementaryAgreement.objects.create(contract=self.contract, name=name)

        # 添加委托人
        if principals:
            for client in principals:
                SupplementaryAgreementParty.objects.create(
                    supplementary_agreement=agreement, client=client, role=PartyRole.PRINCIPAL
                )

        # 添加对方当事人
        if opposing:
            for client in opposing:
                SupplementaryAgreementParty.objects.create(
                    supplementary_agreement=agreement, client=client, role=PartyRole.OPPOSING
                )

        return agreement

    def test_generate_supplementary_agreement_success(self):
        """测试成功生成补充协议"""
        # 创建模板
        template, temp_file_path = self._create_test_template()

        try:
            # 创建补充协议（新增李四作为委托人）
            agreement = self._create_supplementary_agreement(
                name="补充协议一", principals=[self.client1, self.client2]  # 张三 + 李四
            )

            # 调用生成服务
            content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, agreement.id)

            # 验证结果
            assert error is None
            assert content is not None
            assert filename is not None

            # 验证文件名格式
            expected_filename_pattern = r"补充协议一（张三与李四一案代理合同）V1_\d{8}\.docx"
            import re

            assert re.match(expected_filename_pattern, filename)

            # 验证内容不为空
            assert len(content) > 0

        finally:
            # 清理临时文件
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    def test_generate_with_new_principals_and_opposing(self):
        """测试生成包含新增委托人和对方当事人的补充协议"""
        # 创建模板
        template, temp_file_path = self._create_test_template()

        try:
            # 创建补充协议（新增李四作为委托人，王五公司作为对方当事人）
            agreement = self._create_supplementary_agreement(
                name="补充协议二",
                principals=[self.client1, self.client2],  # 张三 + 李四
                opposing=[self.client3],  # 王五公司
            )

            # 调用生成服务
            content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, agreement.id)

            # 验证结果
            assert error is None
            assert content is not None
            assert filename is not None

            # 验证文件名包含正确的补充协议名称
            assert "补充协议二" in filename

        finally:
            # 清理临时文件
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    def test_generate_with_no_new_principals(self):
        """测试生成没有新增委托人的补充协议"""
        # 创建模板
        template, temp_file_path = self._create_test_template()

        try:
            # 创建补充协议（只有原有的张三）
            agreement = self._create_supplementary_agreement(name="补充协议三", principals=[self.client1])  # 只有张三

            # 调用生成服务
            content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, agreement.id)

            # 验证结果
            assert error is None
            assert content is not None
            assert filename is not None

        finally:
            # 清理临时文件
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    def test_contract_not_found_error(self):
        """测试合同不存在的错误处理"""
        # 调用生成服务（使用不存在的合同ID）
        content, filename, error = self.service.generate_supplementary_agreement(99999, 1)

        # 验证错误处理
        assert content is None
        assert filename is None
        assert error == "合同不存在"

    def test_supplementary_agreement_not_found_error(self):
        """测试补充协议不存在的错误处理"""
        # 调用生成服务（使用不存在的补充协议ID）
        content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, 99999)

        # 验证错误处理
        assert content is None
        assert filename is None
        assert error == "补充协议不存在"

    def test_template_not_found_error(self):
        """测试模板不存在的错误处理"""
        # 创建补充协议
        agreement = self._create_supplementary_agreement(name="测试补充协议", principals=[self.client1])

        # 调用生成服务（没有创建模板）
        content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, agreement.id)

        # 验证错误处理
        assert content is None
        assert filename is None
        assert error == "请先添加补充协议模板"

    def test_template_file_not_exists_error(self):
        """测试模板文件不存在的错误处理"""
        # 创建模板记录但不创建实际文件
        template = DocumentTemplateFactory(
            name="不存在的模板",
            template_type=DocumentTemplateType.CONTRACT,
            contract_sub_type=DocumentContractSubType.SUPPLEMENTARY_AGREEMENT,
            contract_types=["civil"],
            is_active=True,
            file_path="/nonexistent/path/template.docx",
        )

        # 创建补充协议
        agreement = self._create_supplementary_agreement(name="测试补充协议", principals=[self.client1])

        # 调用生成服务
        content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, agreement.id)

        # 验证错误处理
        assert content is None
        assert filename is None
        assert error == "模板文件不存在"

    def test_template_matching_logic(self):
        """测试模板匹配逻辑"""
        # 创建多个模板
        template1, temp_file1 = self._create_test_template("civil")
        template2, temp_file2 = self._create_test_template("criminal")
        template3, temp_file3 = self._create_test_template("all")

        try:
            # 测试民事案件匹配 - 应该匹配第一个创建的民事模板
            found_template = self.service.find_supplementary_agreement_template("civil")
            assert found_template is not None
            # 由于可能有多个匹配的模板，我们只验证找到了模板且类型正确
            assert found_template.contract_types == ["civil"] or "all" in found_template.contract_types

            # 测试刑事案件匹配
            found_template = self.service.find_supplementary_agreement_template("criminal")
            assert found_template is not None
            assert found_template.contract_types == ["criminal"] or "all" in found_template.contract_types

            # 测试通用模板匹配（对于没有专门模板的案件类型）
            found_template = self.service.find_supplementary_agreement_template("administrative")
            assert found_template is not None
            assert "all" in found_template.contract_types

        finally:
            # 清理临时文件
            for temp_file in [temp_file1, temp_file2, temp_file3]:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)

    def test_filename_generation(self):
        """测试文件名生成逻辑"""
        # 创建补充协议
        agreement = self._create_supplementary_agreement(name="测试补充协议名称", principals=[self.client1])

        # 测试文件名生成
        filename = self.service.generate_filename(self.contract, agreement)

        # 验证文件名格式
        expected_pattern = r"测试补充协议名称（张三与李四一案代理合同）V1_\d{8}\.docx"
        import re

        assert re.match(expected_pattern, filename)

        # 验证包含正确的组成部分
        assert "测试补充协议名称" in filename
        assert "张三与李四一案代理合同" in filename
        assert "V1_" in filename
        assert ".docx" in filename

    def test_filename_generation_with_empty_names(self):
        """测试空名称时的文件名生成"""
        # 创建没有名称的补充协议
        agreement = SupplementaryAgreement.objects.create(contract=self.contract, name="")  # 空名称

        # 创建没有名称的合同
        contract_no_name = ContractFactory(name="")

        # 测试补充协议名称为空
        filename1 = self.service.generate_filename(self.contract, agreement)
        assert "补充协议（" in filename1  # 使用默认名称

        # 测试合同名称为空
        filename2 = self.service.generate_filename(contract_no_name, agreement)
        assert "（未命名合同）" in filename2  # 使用默认名称

    def test_context_building(self):
        """测试上下文构建逻辑"""
        # 创建补充协议
        agreement = self._create_supplementary_agreement(
            name="测试补充协议", principals=[self.client1, self.client2], opposing=[self.client3]
        )

        # 构建上下文
        context = self.service.build_context(self.contract, agreement)

        # 验证上下文包含必要的数据
        assert isinstance(context, dict)

        # 验证基础占位符
        assert "补充协议名称" in context
        assert "年份" in context

        # 验证委托人信息
        assert "委托人信息" in context
        assert "委托人签名盖章信息" in context

        # 验证条款信息
        assert "委托人主体信息条款" in context
        assert "对方当事人主体信息条款" in context

    def test_principal_data_extraction(self):
        """测试委托人数据提取"""
        # 创建补充协议
        agreement = self._create_supplementary_agreement(name="测试补充协议", principals=[self.client1, self.client2])

        # 测试获取补充协议委托人
        agreement_principals = self.service._get_agreement_principals(agreement)
        assert len(agreement_principals) == 2
        assert self.client1 in agreement_principals
        assert self.client2 in agreement_principals

        # 测试获取原合同委托人
        contract_principals = self.service._get_contract_principals(self.contract)
        assert len(contract_principals) == 1
        assert self.client1 in contract_principals

    def test_opposing_data_extraction(self):
        """测试对方当事人数据提取"""
        # 创建补充协议
        agreement = self._create_supplementary_agreement(
            name="测试补充协议", principals=[self.client1], opposing=[self.client3]
        )

        # 测试获取对方当事人
        opposing_parties = self.service._get_agreement_opposing(agreement)
        assert len(opposing_parties) == 1
        assert self.client3 in opposing_parties

    def test_integration_with_placeholder_services(self):
        """测试与占位符服务的集成"""
        # 创建模板
        template, temp_file_path = self._create_test_template()

        try:
            # 创建复杂的补充协议（包含新增委托人和对方当事人）
            agreement = self._create_supplementary_agreement(
                name="复杂补充协议",
                principals=[self.client1, self.client2],  # 原有 + 新增
                opposing=[self.client3],  # 对方当事人
            )

            # 调用生成服务
            content, filename, error = self.service.generate_supplementary_agreement(self.contract.id, agreement.id)

            # 验证成功生成
            assert error is None
            assert content is not None
            assert filename is not None

            # 验证占位符服务被正确调用（通过检查上下文构建）
            context = self.service.build_context(self.contract, agreement)

            # 验证所有必要的占位符都存在
            required_placeholders = [
                "补充协议名称",
                "年份",
                "委托人信息",
                "委托人签名盖章信息",
                "委托人主体信息条款",
                "对方当事人主体信息条款",
            ]

            for placeholder in required_placeholders:
                assert placeholder in context, f"缺少占位符: {placeholder}"

        finally:
            # 清理临时文件
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
