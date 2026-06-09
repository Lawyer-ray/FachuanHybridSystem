"""
Tests for apps.automation.services.document.document_processing
"""

from __future__ import annotations

from unittest.mock import MagicMock, patch
from pathlib import Path

import pytest


class TestDocumentProcessing:
    """document_processing 模块测试"""

    def test_get_doc_config_default(self) -> None:
        from apps.automation.services.document.document_processing import get_doc_config

        config = get_doc_config()
        assert "DEFAULT_TEXT_LIMIT" in config
        assert "MAX_TEXT_LIMIT" in config
        assert config["DEFAULT_TEXT_LIMIT"] > 0

    def test_apply_pdf_limits_defaults(self) -> None:
        from apps.automation.services.document.document_processing import _apply_pdf_limits

        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(None, None, config)
        assert lim == 1500
        assert page == 1

    def test_apply_pdf_limits_custom(self) -> None:
        from apps.automation.services.document.document_processing import _apply_pdf_limits

        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(5000, 3, config)
        assert lim == 5000
        assert page == 3

    def test_apply_pdf_limits_clamped(self) -> None:
        from apps.automation.services.document.document_processing import _apply_pdf_limits

        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(20000, 10, config)
        assert lim == 10000
        assert page == 5

    def test_extract_document_content_unsupported_format(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_document_content

        test_file = tmp_path / "test.xyz"
        test_file.write_text("content")
        with pytest.raises(ValueError, match="不支持的文件类型"):
            extract_document_content(str(test_file))

    def test_document_extraction_dataclass(self) -> None:
        from apps.automation.services.document.document_processing import DocumentExtraction

        ext = DocumentExtraction(file_path="/test.pdf", text="hello", image_url=None, kind="pdf")
        assert ext.file_path == "/test.pdf"
        assert ext.text == "hello"
        assert ext.kind == "pdf"

    def test_extract_docx_text(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_docx_text

        # Create a minimal docx
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("测试内容第一段")
            doc.add_paragraph("测试内容第二段")
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))
            text = extract_docx_text(str(docx_path), limit=100)
            assert "测试内容" in text
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_extract_docx_text_with_limit(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_docx_text

        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("A" * 100)
            doc.add_paragraph("B" * 100)
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))
            text = extract_docx_text(str(docx_path), limit=50)
            assert len(text) <= 50
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_extract_pdf_text(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_pdf_text

        try:
            import fitz
            pdf_path = tmp_path / "test.pdf"
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Hello World PDF测试")
            doc.save(str(pdf_path))
            doc.close()
            text = extract_pdf_text(str(pdf_path), limit=500)
            assert "Hello" in text or "PDF" in text
        except ImportError:
            pytest.skip("PyMuPDF not installed")


# ============================================================
# DocumentRenamer rename/filename 测试
# ============================================================


class TestDocumentRenamerFilename:
    """DocumentRenamer 生成文件名测试"""

    def test_generate_filename_long_title_truncated(self) -> None:
        from apps.automation.services.sms.document_renamer import DocumentRenamer

        renamer = DocumentRenamer.__new__(DocumentRenamer)
        renamer.title_extraction_limit = 50
        long_title = "这是一个非常非常非常非常非常非常非常长的文书标题"
        filename = renamer.generate_filename(long_title, "案件名", __import__("datetime").date(2025, 1, 1))
        assert filename.endswith(".pdf")

    def test_generate_filename_special_chars(self) -> None:
        from apps.automation.services.sms.document_renamer import DocumentRenamer

        renamer = DocumentRenamer.__new__(DocumentRenamer)
        renamer.title_extraction_limit = 50
        filename = renamer.generate_filename("判决<>书", "张三:*案件", __import__("datetime").date(2025, 6, 1))
        assert "<" not in filename
        assert ">" not in filename
        assert ":" not in filename

    def test_rename_with_fallback_file_not_found(self, tmp_path) -> None:
        from apps.automation.services.sms.document_renamer import DocumentRenamer
        from apps.core.exceptions import ValidationException

        renamer = DocumentRenamer.__new__(DocumentRenamer)
        renamer.title_extraction_limit = 50
        with pytest.raises(ValidationException):
            renamer.rename(str(tmp_path / "nonexistent.pdf"), "案件名", __import__("datetime").date(2025, 1, 1))
