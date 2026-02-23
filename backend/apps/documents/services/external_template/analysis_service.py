"""
外部模板分析服务

负责模板上传校验、文件保存、版本管理。
结构提取和 LLM 分析方法将在后续任务 (6.2, 6.3) 中实现。

Requirements: 1.1, 1.2, 1.3, 1.6, 1.7, 1.8, 9.1, 9.2, 9.3, 9.4
"""

from __future__ import annotations

import logging
import re
import uuid
from pathlib import Path
from typing import TYPE_CHECKING, Any, ClassVar
from xml.etree import ElementTree as ET

from django.conf import settings
from django.core.exceptions import ValidationError
from django.db import transaction
from django.utils.translation import gettext_lazy as _

from .fingerprint_service import FingerprintService

if TYPE_CHECKING:
    from django.core.files.uploadedfile import UploadedFile

    from apps.documents.models.external_template import ExternalTemplate
    from apps.documents.services.placeholders.registry import PlaceholderRegistry

logger: logging.Logger = logging.getLogger(__name__)


class AnalysisService:
    """外部模板分析服务：上传校验 + 结构提取 + LLM 字段映射"""

    MAX_FILE_SIZE: ClassVar[int] = 20 * 1024 * 1024  # 20MB

    def __init__(
        self,
        fingerprint_service: FingerprintService,
        llm_service: Any,
        placeholder_registry: PlaceholderRegistry,
    ) -> None:
        self._fingerprint_service = fingerprint_service
        self._llm_service = llm_service
        self._placeholder_registry = placeholder_registry

    # ------------------------------------------------------------------
    # 上传与校验
    # ------------------------------------------------------------------

    def upload_template(
        self,
        file: UploadedFile,
        name: str,
        category: str,
        source_type: str,
        court_id: int | None,
        organization_name: str,
        uploaded_by: Any,
    ) -> ExternalTemplate:
        """
        上传外部模板：
        1. 校验 .docx 格式、文件大小 ≤ 20MB
        2. 保存文件（UUID 重命名）
        3. 验证 python-docx 可解析
        4. 处理版本管理（同法院/机构 + 类别组合自增版本号）
        5. 创建 ExternalTemplate 记录
        """
        from apps.documents.models.external_template import ExternalTemplate

        self._validate_file(file)

        law_firm_id: int = uploaded_by.law_firm_id
        abs_path, rel_path = self._save_file(file, law_firm_id)

        try:
            self._validate_parseable(abs_path)
        except ValidationError:
            # 解析失败时删除已保存的文件
            if abs_path.exists():
                abs_path.unlink()
                logger.info("已删除无法解析的文件: %s", abs_path)
            raise

        file_size: int = file.size if file.size is not None else 0
        original_filename: str = file.name if file.name else ""

        with transaction.atomic():
            version, deactivated = self._handle_versioning(
                law_firm_id=law_firm_id,
                court_id=court_id,
                category=category,
                source_type=source_type,
                organization_name=organization_name,
            )

            template: ExternalTemplate = ExternalTemplate.objects.create(
                name=name,
                category=category,
                source_type=source_type,
                court_id=court_id,
                organization_name=organization_name,
                file_path=rel_path,
                original_filename=original_filename,
                file_size=file_size,
                version=version,
                is_active=True,
                uploaded_by=uploaded_by,
                law_firm_id=law_firm_id,
            )

        logger.info(
            "模板上传成功: id=%d, name=%s, version=%d, deactivated=%d",
            template.id,
            template.name,
            version,
            deactivated,
        )
        return template

    def _validate_file(self, file: UploadedFile) -> None:
        """校验 .docx 格式和文件大小"""
        filename: str = file.name if file.name else ""
        if not filename.lower().endswith(".docx"):
            logger.info("文件格式校验失败: %s", filename)
            raise ValidationError(_("仅支持 .docx 格式"))

        file_size: int = file.size if file.size is not None else 0
        if file_size > self.MAX_FILE_SIZE:
            logger.info(
                "文件大小超出限制: %s, size=%d",
                filename,
                file_size,
            )
            raise ValidationError(_("文件大小超出限制"))

    def _validate_parseable(self, file_path: Path) -> None:
        """尝试 python-docx 打开文件，验证可解析性"""
        try:
            from docx import Document as DocxDocument

            DocxDocument(str(file_path))
        except Exception as exc:
            logger.info(
                "文件无法解析: %s, error=%s",
                file_path.name,
                str(exc),
            )
            raise ValidationError(
                _("文件无法解析，请检查文件是否损坏或加密")
            ) from exc

    def _save_file(
        self, file: UploadedFile, law_firm_id: int
    ) -> tuple[Path, str]:
        """
        UUID 重命名保存文件

        Returns:
            (绝对路径, 相对于 MEDIA_ROOT 的路径)
        """
        media_root = Path(settings.MEDIA_ROOT)
        rel_dir = Path("documents") / "external_templates" / str(law_firm_id)
        abs_dir = media_root / rel_dir
        abs_dir.mkdir(parents=True, exist_ok=True)

        new_filename = f"{uuid.uuid4()}.docx"
        rel_path = rel_dir / new_filename
        abs_path = media_root / rel_path

        with abs_path.open("wb") as dest:
            for chunk in file.chunks():
                dest.write(chunk)

        logger.info("文件已保存: %s", abs_path)
        return abs_path, str(rel_path)

    # ------------------------------------------------------------------
    # 版本管理
    # ------------------------------------------------------------------

    def _handle_versioning(
        self,
        *,
        law_firm_id: int,
        court_id: int | None,
        category: str,
        source_type: str,
        organization_name: str,
    ) -> tuple[int, int]:
        """
        处理版本管理：同一来源 + 类别组合自增版本号，旧版本 is_active=False

        Returns:
            (新版本号, 被停用的旧版本数量)
        """
        from apps.documents.models.external_template import ExternalTemplate

        existing = ExternalTemplate.objects.filter(
            law_firm_id=law_firm_id,
            category=category,
        )

        if court_id is not None:
            existing = existing.filter(court_id=court_id)
        else:
            existing = existing.filter(
                court__isnull=True,
                organization_name=organization_name,
            )

        max_version: int | None = existing.order_by("-version").values_list(
            "version", flat=True
        ).first()

        new_version: int = (max_version or 0) + 1

        deactivated: int = existing.filter(is_active=True).update(
            is_active=False
        )

        return new_version, deactivated

    # ------------------------------------------------------------------
    # 结构提取 (Requirements: 2.1, 2.2, 2.3, 2.4, 2.5, 2.6, 2.7, 2.8)
    # ------------------------------------------------------------------

    def extract_structure(self, template_id: int) -> dict[str, Any]:
        """
        从 .docx 提取文档结构：段落、表格（含嵌套）、复选框、删除不适用项。

        1. 从数据库加载模板记录，获取 file_path
        2. 使用 python-docx 打开文件
        3. 调用 _extract_paragraphs / _extract_tables / _extract_checkboxes
        4. 将结构 JSON 保存到模板记录的 structure_json 字段
        5. 返回结构 dict
        """
        from apps.documents.models.external_template import ExternalTemplate

        template: ExternalTemplate = ExternalTemplate.objects.get(pk=template_id)
        abs_path = Path(settings.MEDIA_ROOT) / template.file_path

        from docx import Document as DocxDocument

        doc: Any = DocxDocument(str(abs_path))

        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        checkboxes = self._extract_checkboxes(doc)

        structure: dict[str, Any] = {
            "paragraphs": paragraphs,
            "tables": tables,
            "checkboxes": checkboxes,
        }

        template.structure_json = structure
        template.save(update_fields=["structure_json", "updated_at"])

        logger.info(
            "结构提取完成: template_id=%d, paragraphs=%d, tables=%d, checkboxes=%d",
            template_id,
            len(paragraphs),
            len(tables),
            len(checkboxes),
        )
        return structure

    def _extract_paragraphs(self, doc: Any) -> list[dict[str, Any]]:
        """
        提取段落文本和位置索引。

        对每个段落记录：paragraph_index、text、position_locator。
        同时检测"删除不适用项"模式。
        """
        paragraphs: list[dict[str, Any]] = []

        for idx, para in enumerate(doc.paragraphs):
            text: str = para.text.strip()
            if not text:
                continue

            entry: dict[str, Any] = {
                "paragraph_index": idx,
                "text": text,
                "position_locator": {
                    "type": "paragraph",
                    "paragraph_index": idx,
                },
            }

            options = self._detect_delete_inapplicable(text)
            if options is not None:
                entry["delete_inapplicable"] = options
                entry["position_locator"]["type"] = "delete_inapplicable"

            paragraphs.append(entry)

        return paragraphs

    def _extract_tables(self, doc: Any) -> list[dict[str, Any]]:
        """
        提取表格结构（递归嵌套、合并单元格）。

        对每个表格遍历行和单元格，处理合并单元格（通过 _tc 检测），
        递归处理嵌套表格。
        """
        tables: list[dict[str, Any]] = []

        for table_idx, table in enumerate(doc.tables):
            table_data = self._extract_single_table(table, table_idx, [])
            tables.append(table_data)

        return tables

    def _extract_single_table(
        self,
        table: Any,
        table_index: int,
        nested_path: list[int],
    ) -> dict[str, Any]:
        """递归提取单个表格的结构。"""
        rows_data: list[dict[str, Any]] = []
        seen_cells: set[int] = set()

        for row_idx, row in enumerate(table.rows):
            cells_data: list[dict[str, Any]] = []

            for col_idx, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                is_merged = cell_id in seen_cells
                seen_cells.add(cell_id)

                if is_merged:
                    continue

                text: str = cell.text.strip()
                cell_entry: dict[str, Any] = {
                    "row": row_idx,
                    "col": col_idx,
                    "text": text,
                    "position_locator": {
                        "type": "table_cell",
                        "table_index": table_index,
                        "row": row_idx,
                        "col": col_idx,
                        "nested_table_path": nested_path,
                    },
                }

                options = self._detect_delete_inapplicable(text)
                if options is not None:
                    cell_entry["delete_inapplicable"] = options
                    cell_entry["position_locator"]["type"] = "delete_inapplicable"

                # 递归处理嵌套表格
                nested_tables: list[dict[str, Any]] = []
                for nested_idx, nested_table in enumerate(cell.tables):
                    nested_data = self._extract_single_table(
                        nested_table,
                        nested_idx,
                        [*nested_path, table_index],
                    )
                    nested_tables.append(nested_data)

                if nested_tables:
                    cell_entry["nested_tables"] = nested_tables

                cells_data.append(cell_entry)

            rows_data.append({"row_index": row_idx, "cells": cells_data})

        return {
            "table_index": table_index,
            "nested_table_path": nested_path,
            "rows": rows_data,
        }

    def _extract_checkboxes(self, doc: Any) -> list[dict[str, Any]]:
        """
        识别复选框控件。

        解析文档 XML，查找 w14:checkbox 或 w:sdt 中包含复选框的结构化文档标签。
        """
        checkboxes: list[dict[str, Any]] = []
        ns: dict[str, str] = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        }

        body_xml: str = doc.element.xml
        try:
            root: ET.Element = ET.fromstring(body_xml)
        except ET.ParseError:
            logger.info("复选框提取: XML 解析失败，跳过")
            return checkboxes

        # 查找所有 sdt（结构化文档标签）
        sdt_elements = root.iter(f"{{{ns['w']}}}sdt")
        checkbox_index = 0

        for sdt in sdt_elements:
            sdt_pr = sdt.find(f"{{{ns['w']}}}sdtPr", ns)
            if sdt_pr is None:
                continue

            # 检查是否为复选框：w14:checkbox 或 w:checkbox
            is_checkbox = (
                sdt_pr.find(f"{{{ns['w14']}}}checkbox", ns) is not None
                or sdt_pr.find(f"{{{ns['w']}}}checkbox", ns) is not None
            )
            if not is_checkbox:
                continue

            # 提取关联文本标签
            sdt_content = sdt.find(f"{{{ns['w']}}}sdtContent", ns)
            label_text = ""
            if sdt_content is not None:
                runs = sdt_content.iter(f"{{{ns['w']}}}t")
                label_parts: list[str] = [
                    r.text for r in runs if r.text
                ]
                label_text = "".join(label_parts)

            # 尝试获取 checked 状态
            checked = False
            checkbox_el = sdt_pr.find(f"{{{ns['w14']}}}checkbox", ns)
            if checkbox_el is None:
                checkbox_el = sdt_pr.find(f"{{{ns['w']}}}checkbox", ns)
            if checkbox_el is not None:
                checked_el = checkbox_el.find(
                    f"{{{ns['w14']}}}checked", ns
                )
                if checked_el is None:
                    checked_el = checkbox_el.find(
                        f"{{{ns['w']}}}checked", ns
                    )
                if checked_el is not None:
                    val = checked_el.get(
                        f"{{{ns['w14']}}}val",
                        checked_el.get(f"{{{ns['w']}}}val", "0"),
                    )
                    checked = val in ("1", "true")

            checkboxes.append({
                "checkbox_index": checkbox_index,
                "label": label_text,
                "checked": checked,
                "position_locator": {
                    "type": "checkbox",
                    "checkbox_index": checkbox_index,
                },
            })
            checkbox_index += 1

        logger.info("复选框提取完成: 共 %d 个", len(checkboxes))
        return checkboxes

    def _detect_delete_inapplicable(self, text: str) -> list[str] | None:
        """
        检测"删除不适用项"格式（如"自然人/法人/非法人组织"）。

        使用正则匹配以 / 或 ／ 分隔的选项模式。
        返回选项列表，若未检测到则返回 None。
        """
        # 匹配 "A/B/C" 或 "A／B／C" 模式，至少 2 个选项
        pattern = r"([\u4e00-\u9fff\w]+(?:[/／][\u4e00-\u9fff\w]+){1,})"
        match = re.search(pattern, text)
        if match is None:
            return None

        matched_text: str = match.group(1)
        options: list[str] = re.split(r"[/／]", matched_text)

        # 过滤：每个选项至少 1 个中文字符，才认为是"删除不适用项"
        has_chinese = all(
            re.search(r"[\u4e00-\u9fff]", opt) for opt in options
        )
        if not has_chinese:
            return None

        return options
