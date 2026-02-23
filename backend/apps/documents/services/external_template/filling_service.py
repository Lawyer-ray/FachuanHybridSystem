"""
外部模板填充服务

负责占位符取值、填充预览、自定义字段获取、单次模板填充。
批量填充方法将在后续任务 (7.3) 中实现。

Requirements: 6.1-6.10, 10.1-10.6, 11.3, 15.3, 15.4, 15.6,
              16.1-16.6, 17.1-17.3
"""

from __future__ import annotations

import logging
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any
from xml.etree import ElementTree as ET

from django.conf import settings
from django.utils.translation import gettext_lazy as _

if TYPE_CHECKING:
    from apps.documents.services.placeholders.registry import PlaceholderRegistry

logger: logging.Logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class FillPreviewItem:
    """填充预览项"""

    position_description: str
    semantic_label: str
    fill_value: str
    value_source: str  # "auto" | "manual" | "empty"
    fill_type: str
    mapping_id: int


@dataclass(frozen=True)
class FillReport:
    """填充报告"""

    total_fields: int
    filled_count: int
    skipped_count: int
    manual_needed: list[str]
    errors: list[str]


class FillingService:
    """模板填充服务：占位符取值 + 填充预览 + 自定义字段"""

    def __init__(self, placeholder_registry: PlaceholderRegistry) -> None:
        self._placeholder_registry = placeholder_registry

    # ------------------------------------------------------------------
    # 预览
    # ------------------------------------------------------------------

    def generate_preview(
        self,
        template_id: int,
        case_id: int,
        party_id: int | None = None,
        custom_values: dict[str, str] | None = None,
    ) -> list[FillPreviewItem]:
        """
        生成填充预览：
        1. 获取模板的所有 FieldMapping（按 sort_order 排序）
        2. 从占位符体系获取案件数据（含当事人数据）
        3. 合并自定义值
        4. 返回每个字段的预览信息（位置、语义、值、来源）
        """
        from apps.documents.models.external_template import (
            ExternalTemplateFieldMapping,
        )

        mappings = ExternalTemplateFieldMapping.objects.filter(
            template_id=template_id,
        ).order_by("sort_order", "id")

        placeholder_values: dict[str, str] = self._get_placeholder_values(
            case_id, party_id
        )
        merged_custom: dict[str, str] = custom_values or {}

        preview_items: list[FillPreviewItem] = []
        for mapping in mappings:
            fill_value: str = ""
            value_source: str = "empty"

            if mapping.placeholder_key and mapping.placeholder_key in placeholder_values:
                fill_value = str(placeholder_values[mapping.placeholder_key])
                value_source = "auto"
            elif not mapping.placeholder_key and mapping.semantic_label in merged_custom:
                fill_value = merged_custom[mapping.semantic_label]
                value_source = "manual"

            preview_items.append(
                FillPreviewItem(
                    position_description=mapping.position_description,
                    semantic_label=mapping.semantic_label,
                    fill_value=fill_value,
                    value_source=value_source,
                    fill_type=mapping.fill_type,
                    mapping_id=mapping.id,
                )
            )

        logger.info(
            "填充预览生成: template_id=%d, case_id=%d, items=%d",
            template_id,
            case_id,
            len(preview_items),
        )
        return preview_items

    def get_custom_fields(self, template_id: int) -> list[dict[str, Any]]:
        """获取需要手动输入的自定义字段列表（placeholder_key 为空的映射）"""
        from apps.documents.models.external_template import (
            ExternalTemplateFieldMapping,
        )

        mappings = ExternalTemplateFieldMapping.objects.filter(
            template_id=template_id,
            placeholder_key="",
        ).order_by("sort_order", "id")

        fields: list[dict[str, Any]] = []
        for mapping in mappings:
            fields.append({
                "mapping_id": mapping.id,
                "semantic_label": mapping.semantic_label,
                "fill_type": mapping.fill_type,
                "options": mapping.options,
                "position_description": mapping.position_description,
            })

        logger.info(
            "自定义字段获取: template_id=%d, count=%d",
            template_id,
            len(fields),
        )
        return fields

    # ------------------------------------------------------------------
    # 内部辅助
    # ------------------------------------------------------------------

    def _get_placeholder_values(
        self, case_id: int, party_id: int | None = None
    ) -> dict[str, str]:
        """
        从占位符体系获取案件+当事人的所有占位符值。

        1. 构建 context_data（case_id, party_id）
        2. 遍历 registry 中所有服务，调用 generate(context_data)
        3. 合并所有结果为 dict[str, str]
        """
        context_data: dict[str, Any] = {"case_id": case_id}
        if party_id is not None:
            context_data["party_id"] = party_id

        all_values: dict[str, str] = {}
        services = self._placeholder_registry.get_all_services()

        for service in services:
            try:
                result: dict[str, Any] = service.generate(context_data)
                for key, value in result.items():
                    all_values[key] = str(value) if value is not None else ""
            except Exception:
                logger.exception(
                    "占位符服务 %s 生成失败: case_id=%d",
                    service.name,
                    case_id,
                )

        logger.info(
            "占位符值获取: case_id=%d, party_id=%s, keys=%d",
            case_id,
            party_id,
            len(all_values),
        )
        return all_values

    # ------------------------------------------------------------------
    # 单次填充
    # ------------------------------------------------------------------

    def fill_template(
        self,
        template_id: int,
        case_id: int,
        party_id: int | None = None,
        custom_values: dict[str, str] | None = None,
        filled_by: Any = None,
    ) -> Any:
        """
        执行单次填充：
        1. 获取占位符值 + 自定义值
        2. 打开模板 .docx 副本
        3. 按 FieldMapping 逐一写入值
        4. 保存生成文件
        5. 创建 FillRecord
        6. 返回 FillRecord

        Requirements: 6.1, 6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 6.8, 6.9, 6.10,
                      10.1, 10.2, 10.3, 10.4, 10.5, 10.6, 11.3, 15.3, 15.4,
                      15.6, 16.3, 16.4, 16.6
        """
        from docx import Document

        from apps.documents.models.external_template import (
            ExternalTemplate,
            ExternalTemplateFieldMapping,
        )
        from apps.documents.models.fill_record import FillRecord

        template: ExternalTemplate = ExternalTemplate.objects.get(id=template_id)
        mappings = ExternalTemplateFieldMapping.objects.filter(
            template_id=template_id,
        ).order_by("sort_order", "id")

        # 1. 获取占位符值 + 合并自定义值
        placeholder_values: dict[str, str] = self._get_placeholder_values(
            case_id, party_id
        )
        merged_custom: dict[str, str] = custom_values or {}

        # 2. 打开模板 .docx 副本
        template_path: Path = Path(settings.MEDIA_ROOT) / template.file_path
        doc: Document = Document(str(template_path))

        # 3. 逐一写入
        filled_count: int = 0
        skipped_count: int = 0
        manual_needed: list[str] = []
        errors: list[str] = []

        for mapping in mappings:
            value: str = ""
            # 确定值来源
            if mapping.placeholder_key and mapping.placeholder_key in placeholder_values:
                value = placeholder_values[mapping.placeholder_key]
            elif not mapping.placeholder_key and mapping.semantic_label in merged_custom:
                value = merged_custom[mapping.semantic_label]
            elif not mapping.placeholder_key:
                # 无占位符键且用户未提供自定义值 → 跳过
                manual_needed.append(mapping.semantic_label)
                skipped_count += 1
                logger.info(
                    "跳过需手动填写字段: template_id=%d, label=%s",
                    template_id,
                    mapping.semantic_label,
                )
                continue
            elif mapping.placeholder_key not in placeholder_values:
                # 有占位符键但无对应值 → 跳过
                skipped_count += 1
                logger.info(
                    "占位符值缺失: template_id=%d, key=%s",
                    template_id,
                    mapping.placeholder_key,
                )
                continue

            # 根据 fill_type 调用对应写入方法
            success: bool = False
            try:
                if mapping.fill_type == "text":
                    success = self._write_text(doc, mapping.position_locator, value)
                elif mapping.fill_type == "checkbox":
                    success = self._write_checkbox(doc, mapping.position_locator, value)
                elif mapping.fill_type == "delete_inapplicable":
                    success = self._write_delete_inapplicable(
                        doc, mapping.position_locator, value, mapping.options
                    )
                else:
                    logger.warning(
                        "未知填充类型: template_id=%d, fill_type=%s",
                        template_id,
                        mapping.fill_type,
                    )
                    skipped_count += 1
                    continue
            except Exception:
                logger.exception(
                    "填充写入失败: template_id=%d, locator=%s",
                    template_id,
                    mapping.position_locator,
                )
                errors.append(
                    str(_("位置 %(label)s 写入失败") % {"label": mapping.semantic_label})
                )
                continue

            if success:
                filled_count += 1
            else:
                skipped_count += 1
                errors.append(
                    str(_("位置 %(label)s 写入未成功") % {"label": mapping.semantic_label})
                )

        # 4. 保存生成文件
        output_dir: Path = Path(settings.MEDIA_ROOT) / "documents" / "external_filled" / str(case_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_uuid: str = str(uuid.uuid4())
        output_abs: Path = output_dir / f"{output_uuid}.docx"
        doc.save(str(output_abs))

        output_relative: str = str(
            Path("documents") / "external_filled" / str(case_id) / f"{output_uuid}.docx"
        )

        # 5. 生成输出文件名
        party_name: str | None = None
        if party_id is not None:
            try:
                from apps.cases.models import CaseParty

                party_obj = CaseParty.objects.get(id=party_id)
                party_name = str(party_obj)
            except Exception:
                logger.warning("获取当事人名称失败: party_id=%d", party_id)

        is_confirmed: bool = all(m.is_confirmed for m in mappings)
        output_name: str = self._generate_output_filename(
            template.name, party_name, is_confirmed
        )

        # 6. 构建填充报告
        report: dict[str, Any] = {
            "total_fields": len(list(mappings)),
            "filled_count": filled_count,
            "skipped_count": skipped_count,
            "manual_needed": manual_needed,
            "errors": errors,
        }

        # 7. 创建 FillRecord
        record: FillRecord = FillRecord.objects.create(
            case_id=case_id,
            template=template,
            party_id=party_id,
            filled_by=filled_by,
            file_path=output_relative,
            original_output_name=output_name,
            report_json=report,
            custom_values=merged_custom,
        )

        logger.info(
            "填充完成: record_id=%d, template_id=%d, case_id=%d, "
            "filled=%d, skipped=%d, errors=%d",
            record.id,
            template_id,
            case_id,
            filled_count,
            skipped_count,
            len(errors),
        )
        return record

    # ------------------------------------------------------------------
    # 写入方法
    # ------------------------------------------------------------------

    def _write_text(
        self, doc: Any, locator: dict[str, Any], value: str
    ) -> bool:
        """
        写入文本值，保留原有格式属性。

        支持两种定位方式：
        - paragraph: 通过 paragraph_index 定位段落
        - table_cell: 通过 table_index + row + col 定位表格单元格

        Requirements: 6.2, 6.3, 10.1, 10.2, 10.3
        """
        locator_type: str = locator.get("type", "")

        try:
            if locator_type == "paragraph":
                para_index: int = locator.get("paragraph_index", 0)
                if para_index >= len(doc.paragraphs):
                    logger.warning(
                        "段落索引越界: index=%d, total=%d",
                        para_index,
                        len(doc.paragraphs),
                    )
                    return False

                paragraph = doc.paragraphs[para_index]
                runs = paragraph.runs
                if not runs:
                    # 无 run 时直接添加
                    paragraph.add_run(value)
                    return True

                # 保留第一个 run 的格式，替换文本
                first_run = runs[0]
                first_run.text = value
                # 清除后续 run 的文本
                for run in runs[1:]:
                    run.text = ""
                return True

            elif locator_type == "table_cell":
                table_index: int = locator.get("table_index", 0)
                row: int = locator.get("row", 0)
                col: int = locator.get("col", 0)

                if table_index >= len(doc.tables):
                    logger.warning(
                        "表格索引越界: index=%d, total=%d",
                        table_index,
                        len(doc.tables),
                    )
                    return False

                table = doc.tables[table_index]
                if row >= len(table.rows) or col >= len(table.columns):
                    logger.warning(
                        "单元格索引越界: row=%d, col=%d, "
                        "rows=%d, cols=%d",
                        row,
                        col,
                        len(table.rows),
                        len(table.columns),
                    )
                    return False

                cell = table.cell(row, col)
                # 写入单元格第一个段落
                if cell.paragraphs:
                    paragraph = cell.paragraphs[0]
                    runs = paragraph.runs
                    if not runs:
                        paragraph.add_run(value)
                    else:
                        runs[0].text = value
                        for run in runs[1:]:
                            run.text = ""
                else:
                    cell.text = value
                return True

            else:
                logger.warning("未知定位器类型: %s", locator_type)
                return False

        except Exception:
            logger.exception("写入文本失败: locator=%s", locator)
            return False

    def _write_checkbox(
        self, doc: Any, locator: dict[str, Any], value: str
    ) -> bool:
        """
        设置复选框勾选状态。

        通过解析文档 XML 找到复选框控件并设置 checked 状态。
        value 为 "true"/"1" 时勾选，否则取消勾选。

        Requirements: 6.2
        """
        try:
            checkbox_index: int = locator.get("checkbox_index", 0)
            checked: bool = value.lower() in ("true", "1", "yes")

            # 解析文档 XML 查找复选框
            body_xml: str = doc.element.xml
            root: ET.Element = ET.fromstring(body_xml)

            # Word 复选框命名空间
            ns: dict[str, str] = {
                "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            }

            # 查找所有 w14:checkbox 元素
            checkboxes: list[ET.Element] = root.findall(".//w14:checkbox", ns)

            if checkbox_index >= len(checkboxes):
                # 尝试旧版复选框格式 (w:fldChar + w:ffData)
                ff_checkboxes: list[ET.Element] = root.findall(
                    ".//w:ffData/w:checkBox", ns
                )
                if checkbox_index >= len(ff_checkboxes):
                    logger.warning(
                        "复选框索引越界: index=%d, w14=%d, ff=%d",
                        checkbox_index,
                        len(checkboxes),
                        len(ff_checkboxes),
                    )
                    return False

                # 旧版复选框：设置 w:default 或 w:checked
                cb: ET.Element = ff_checkboxes[checkbox_index]
                checked_elem: ET.Element | None = cb.find("w:checked", ns)
                if checked_elem is None:
                    checked_elem = cb.find("w:default", ns)

                if checked_elem is not None:
                    checked_elem.set(
                        f"{{{ns['w']}}}val",
                        "1" if checked else "0",
                    )
                return True

            # 新版复选框 (w14:checkbox)
            cb_elem: ET.Element = checkboxes[checkbox_index]
            checked_state: ET.Element | None = cb_elem.find("w14:checked", ns)
            if checked_state is not None:
                checked_state.set(
                    f"{{{ns['w14']}}}val",
                    "1" if checked else "0",
                )
            return True

        except Exception:
            logger.exception("写入复选框失败: locator=%s", locator)
            return False

    def _write_delete_inapplicable(
        self,
        doc: Any,
        locator: dict[str, Any],
        value: str,
        options: list[str],
    ) -> bool:
        """
        删除不适用项：保留匹配项，删除其余选项及分隔符（/）。

        Requirements: 6.2, 6.10
        """
        try:
            locator_type: str = locator.get("type", "")

            paragraph = None
            if locator_type == "paragraph":
                para_index: int = locator.get("paragraph_index", 0)
                if para_index >= len(doc.paragraphs):
                    logger.warning(
                        "段落索引越界: index=%d, total=%d",
                        para_index,
                        len(doc.paragraphs),
                    )
                    return False
                paragraph = doc.paragraphs[para_index]

            elif locator_type == "table_cell":
                table_index: int = locator.get("table_index", 0)
                row: int = locator.get("row", 0)
                col: int = locator.get("col", 0)

                if table_index >= len(doc.tables):
                    return False

                table = doc.tables[table_index]
                if row >= len(table.rows) or col >= len(table.columns):
                    return False

                cell = table.cell(row, col)
                if cell.paragraphs:
                    paragraph = cell.paragraphs[0]

            elif locator_type == "delete_inapplicable":
                # 兼容直接使用 paragraph_index 的定位器
                para_index = locator.get("paragraph_index", 0)
                if para_index < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_index]

            if paragraph is None:
                logger.warning("无法定位删除不适用项段落: locator=%s", locator)
                return False

            # 获取当前文本
            current_text: str = paragraph.text

            # 在选项中找到匹配项
            matched_option: str = value
            if not matched_option and options:
                matched_option = options[0]

            # 构建新文本：用匹配项替换 "选项A/选项B/选项C" 格式
            for option in options:
                if option != matched_option:
                    # 删除不匹配的选项及其前后的分隔符
                    current_text = current_text.replace(f"/{option}", "")
                    current_text = current_text.replace(f"{option}/", "")
                    current_text = current_text.replace(option, "")

            # 清理多余的分隔符
            while "//" in current_text:
                current_text = current_text.replace("//", "/")
            current_text = current_text.strip("/").strip()

            # 写回段落，保留格式
            runs = paragraph.runs
            if runs:
                runs[0].text = current_text
                for run in runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(current_text)

            return True

        except Exception:
            logger.exception("删除不适用项失败: locator=%s", locator)
            return False

    # ------------------------------------------------------------------
    # 文件名生成
    # ------------------------------------------------------------------

    def _generate_output_filename(
        self,
        template_name: str,
        party_name: str | None = None,
        is_confirmed: bool = True,
    ) -> str:
        """
        生成输出文件名。

        格式：
        - 有当事人: "{template_name}_{party_name}.docx"
        - 无当事人: "{template_name}.docx"
        - 未确认映射: "[未确认]{template_name}_{party_name}.docx"

        Requirements: 5.8, 15.4
        """
        base_name: str = template_name
        if party_name:
            base_name = f"{template_name}_{party_name}"

        filename: str = f"{base_name}.docx"

        if not is_confirmed:
            filename = f"[{_('未确认')}]{filename}"

        return filename
