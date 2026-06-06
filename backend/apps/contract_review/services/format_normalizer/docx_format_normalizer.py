"""合同格式规范化工具 - 直接复制修订版格式

通过读取修订版文档的精确格式，逐段落复制到原始文档。
"""

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

logger = logging.getLogger(__name__)

# 页边距（EMU 单位）
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)


class DocxFormatNormalizer:
    """合同格式规范化器 - 直接复制修订版格式"""

    def __init__(
        self, input_path: str | Path, output_path: str | Path | None = None, reference_path: str | Path | None = None
    ):
        self.input_path = Path(input_path)
        self.output_path = (
            Path(output_path)
            if output_path
            else self.input_path.parent / f"{self.input_path.stem}_规范化{self.input_path.suffix}"
        )
        self.reference_path = Path(reference_path) if reference_path else None
        self.doc: Any = None
        self.ref_doc: Any = None

    def normalize(self) -> Path:
        """执行格式规范化，返回输出文件路径"""
        logger.info("开始规范化: %s", self.input_path)
        self.doc = Document(str(self.input_path))

        # 加载参考文档
        if self.reference_path and self.reference_path.exists():
            self.ref_doc = Document(str(self.reference_path))
            logger.info("已加载参考文档: %s", self.reference_path)

        # 1. 设置页边距
        self._normalize_margins()

        # 2. 定义编号样式
        self._setup_numbering()

        # 3. 规范化段落格式
        if self.ref_doc:
            # 使用参考文档的格式
            self._normalize_with_reference()
        else:
            # 使用默认格式
            self._normalize_default()

        # 4. 保存
        assert self.doc is not None, "doc 应在 normalize 开头初始化"
        self.doc.save(str(self.output_path))
        logger.info("规范化完成: %s", self.output_path)
        return self.output_path

    def _normalize_margins(self) -> None:
        """统一页边距为 A4 标准"""
        assert self.doc is not None
        for section in self.doc.sections:
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT
        logger.debug("页边距已标准化")

    def _setup_numbering(self) -> None:
        """设置自动编号样式"""
        assert self.doc is not None
        try:
            numbering_part = self.doc.part.numbering_part
            numbering_elm = numbering_part._element
        except (KeyError, NotImplementedError):
            numbering_elm = self._create_numbering_part()

        # 创建 abstractNum
        abstractNum = OxmlElement("w:abstractNum")
        abstractNum.set(qn("w:abstractNumId"), "0")

        # 一级：一、二、三...
        lvl0 = self._create_level("0", "chineseCounting", "%1、", "400")
        abstractNum.append(lvl0)

        # 二级：1. 2. 3.
        lvl1 = self._create_level("1", "decimal", "%2．", "400")
        abstractNum.append(lvl1)

        # 三级：（1）（2）（3）
        lvl2 = self._create_level("2", "decimal", "（%3）", "402")
        abstractNum.append(lvl2)

        # 插入到 numbering 元素开头
        numbering_elm.insert(0, abstractNum)

        # 创建 num 实例（一级标题）
        num1 = OxmlElement("w:num")
        num1.set(qn("w:numId"), "1")
        abstractNumRef1 = OxmlElement("w:abstractNumId")
        abstractNumRef1.set(qn("w:val"), "0")
        num1.append(abstractNumRef1)
        numbering_elm.append(num1)

        # 创建 num 实例（二级标题）
        num2 = OxmlElement("w:num")
        num2.set(qn("w:numId"), "2")
        abstractNumRef2 = OxmlElement("w:abstractNumId")
        abstractNumRef2.set(qn("w:val"), "0")
        num2.append(abstractNumRef2)
        numbering_elm.append(num2)

        # 创建 num 实例（三级标题）
        num3 = OxmlElement("w:num")
        num3.set(qn("w:numId"), "3")
        abstractNumRef3 = OxmlElement("w:abstractNumId")
        abstractNumRef3.set(qn("w:val"), "0")
        num3.append(abstractNumRef3)
        numbering_elm.append(num3)

        # 如果是新创建的 part，需要更新其内容
        if hasattr(self, "_numbering_part"):
            from lxml import etree

            self._numbering_part._blob = etree.tostring(
                numbering_elm, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        logger.debug("编号样式已创建")

    def _create_numbering_part(self) -> Any:
        """手动创建 numbering part"""
        assert self.doc is not None
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part
        from lxml import etree

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        numbering_elm = etree.Element(qn("w:numbering"), nsmap=nsmap)
        numbering_xml = etree.tostring(numbering_elm, xml_declaration=True, encoding="UTF-8", standalone=True)

        part_name = PackURI("/word/numbering.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"
        numbering_part = Part(part_name, content_type, numbering_xml, self.doc.part.package)
        self.doc.part.relate_to(numbering_part, RT.NUMBERING)
        self._numbering_part = numbering_part

        return numbering_elm

    def _create_level(self, ilvl: str, num_fmt: str, level_text: str, first_line: str) -> Any:
        """创建编号级别定义"""
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), ilvl)
        lvl.set(qn("w:tentative"), "0")

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), num_fmt)
        lvl.append(numFmt)

        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "nothing")
        lvl.append(suff)

        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), level_text)
        lvl.append(lvlText)

        lvlJc = OxmlElement("w:lvlJc")
        lvlJc.set(qn("w:val"), "left")
        lvl.append(lvlJc)

        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:firstLine"), first_line)
        pPr.append(ind)
        lvl.append(pPr)

        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:hint"), "eastAsia")
        rPr.append(rFonts)
        lvl.append(rPr)

        return lvl

    def _normalize_with_reference(self) -> None:
        """使用参考文档的格式进行规范化"""
        assert self.ref_doc is not None
        assert self.doc is not None
        ref_paras = self.ref_doc.paragraphs
        doc_paras = self.doc.paragraphs

        # 逐段落复制格式
        for i, para in enumerate(doc_paras):
            if i < len(ref_paras):
                # 从参考文档获取格式
                ref_para = ref_paras[i]
                self._copy_paragraph_format(ref_para, para, i)
            else:
                # 超出参考文档范围，使用默认格式
                self._apply_default_format(para, i)

        logger.debug("段落格式已规范化（使用参考文档）")

    def _normalize_default(self) -> None:
        """使用默认格式进行规范化"""
        assert self.doc is not None
        for i, para in enumerate(self.doc.paragraphs):
            self._apply_default_format(para, i)

        logger.debug("段落格式已规范化（使用默认格式）")

    def _copy_paragraph_format(self, ref_para: Any, doc_para: Any, index: int) -> None:
        """从参考文档复制段落格式"""
        # 获取参考文档的段落格式
        ref_pPr = ref_para._element.find(qn("w:pPr"))
        doc_pPr = doc_para._element.get_or_add_pPr()

        # 清除旧的格式
        self._clear_old_format(doc_pPr)

        # 复制 pPr 属性
        if ref_pPr is not None:
            # 复制对齐方式
            ref_jc = ref_pPr.find(qn("w:jc"))
            if ref_jc is not None:
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), ref_jc.get(qn("w:val")))
                doc_pPr.append(jc)

            # 复制行距
            ref_spacing = ref_pPr.find(qn("w:spacing"))
            if ref_spacing is not None:
                spacing = OxmlElement("w:spacing")
                for attr in ["w:line", "w:lineRule", "w:before", "w:after"]:
                    val = ref_spacing.get(qn(attr))
                    if val is not None:
                        spacing.set(qn(attr), val)
                doc_pPr.append(spacing)

            # 复制缩进
            ref_ind = ref_pPr.find(qn("w:ind"))
            if ref_ind is not None:
                ind = OxmlElement("w:ind")
                for attr in ["w:left", "w:firstLine", "w:hanging"]:
                    val = ref_ind.get(qn(attr))
                    if val is not None:
                        ind.set(qn(attr), val)
                doc_pPr.append(ind)

            # 复制编号
            ref_numPr = ref_pPr.find(qn("w:numPr"))
            if ref_numPr is not None:
                numPr = OxmlElement("w:numPr")
                ref_ilvl = ref_numPr.find(qn("w:ilvl"))
                if ref_ilvl is not None:
                    ilvl = OxmlElement("w:ilvl")
                    ilvl.set(qn("w:val"), ref_ilvl.get(qn("w:val")))
                    numPr.append(ilvl)
                ref_numId = ref_numPr.find(qn("w:numId"))
                if ref_numId is not None:
                    numId = OxmlElement("w:numId")
                    numId.set(qn("w:val"), ref_numId.get(qn("w:val")))
                    numPr.append(numId)
                doc_pPr.append(numPr)

        # 复制 run 格式
        self._copy_run_format(ref_para, doc_para)

    def _copy_run_format(self, ref_para: Any, doc_para: Any) -> None:
        """从参考文档复制 run 格式"""
        # 获取参考文档的 run 格式
        if ref_para.runs:
            ref_rPr = ref_para.runs[0]._element.find(qn("w:rPr"))
        else:
            ref_rPr = None

        # 如果文档没有 run，但参考文档有 run 格式，需要添加 run
        if not doc_para.runs and ref_rPr is not None:
            # 创建一个空的 run
            run_elem = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            run_elem.append(rPr)
            doc_para._element.append(run_elem)

        # 应用到文档的每个 run
        for run in doc_para.runs:
            rPr = run._element.get_or_add_rPr()

            # 清除旧的格式
            for tag in ["w:rFonts", "w:sz", "w:szCs", "w:b"]:
                old = rPr.find(qn(tag))
                if old is not None:
                    rPr.remove(old)

            # 复制格式
            if ref_rPr is not None:
                # 复制字体
                ref_rFonts = ref_rPr.find(qn("w:rFonts"))
                if ref_rFonts is not None:
                    rFonts = OxmlElement("w:rFonts")
                    for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
                        val = ref_rFonts.get(qn(attr))
                        if val is not None:
                            rFonts.set(qn(attr), val)
                    rPr.insert(0, rFonts)

                # 复制字号
                ref_sz = ref_rPr.find(qn("w:sz"))
                if ref_sz is not None:
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), ref_sz.get(qn("w:val")))
                    rPr.append(sz)

                ref_szCs = ref_rPr.find(qn("w:szCs"))
                if ref_szCs is not None:
                    szCs = OxmlElement("w:szCs")
                    szCs.set(qn("w:val"), ref_szCs.get(qn("w:val")))
                    rPr.append(szCs)

                # 复制加粗
                ref_b = ref_rPr.find(qn("w:b"))
                if ref_b is not None:
                    b = OxmlElement("w:b")
                    rPr.append(b)

    def _apply_default_format(self, para: Any, index: int) -> None:
        """应用默认格式"""
        pPr = para._element.get_or_add_pPr()

        # 清除旧的格式
        self._clear_old_format(pPr)

        # 基础格式：行距 360，左缩进 0
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "360")
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)

        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        pPr.append(ind)

        # 根据段落内容设置编号，并删除手动编号
        text = para.text.strip()
        if text:
            # 先尝试规则方法
            level = self._detect_level_by_rules(text)

            # 如果规则方法无法判断，使用LLM
            if level == -1:
                try:
                    from .llm_helper import ContractStructureAnalyzer
                    analyzer = ContractStructureAnalyzer()
                    result = analyzer.analyze_paragraph_level(text)
                    level = result["level"]
                    logger.debug(f"LLM判断段落 '{text[:30]}...' 的层级为 {level}")
                except Exception as e:
                    logger.warning(f"LLM判断失败: {e}")
                    level = -1

            # 应用编号
            if level == 0:
                self._apply_numbering(para, "1", "0")  # numId=1, ilvl=0
                self._remove_manual_numbering(para, "一级")
            elif level == 1:
                self._apply_numbering(para, "2", "1")  # numId=2, ilvl=0
                self._remove_manual_numbering(para, "二级")
            elif level == 2:
                self._apply_numbering(para, "3", "2")  # numId=3, ilvl=0
                self._remove_manual_numbering(para, "三级")

    def _detect_level_by_rules(self, text: str) -> int:
        """使用规则方法检测段落层级"""
        # 检测一级标题（一、二、三... 或 （一）、（二）、...）
        if (len(text) >= 2 and text[0] in "一二三四五六七八九十" and text[1] == "、"):
            return 0
        elif (len(text) >= 4 and text.startswith("（") and text[1] in "一二三四五六七八九十" and text[2] == "）"):
            return 0
        # 检测二级标题（1. 2. 3.）
        elif len(text) >= 2 and text[0].isdigit() and text[1] == ".":
            return 1
        # 检测三级标题（（1）（2）（3））
        elif len(text) >= 4 and text.startswith("（") and text[1].isdigit() and text[2] == "）":
            return 2

        # 检测常见的一级标题关键词
        level0_keywords = ["服务内容", "服务范围", "费用", "保密义务", "责任限制", "免责条款", "合同期限", "违约责任"]
        for keyword in level0_keywords:
            if keyword in text and len(text) < 20:
                return 0

        return -1  # 无法判断

    def _remove_manual_numbering(self, para: Any, level_type: str) -> None:
        """删除段落中的手动编号文本"""
        text = para.text
        if not text:
            return

        new_text = text

        if level_type == "一级":
            # 删除"（一）、"或"一、"格式
            if new_text.startswith("（") and len(new_text) >= 4 and new_text[2] == "）":
                new_text = new_text[4:]  # 删除"（X）、"
            elif len(new_text) >= 2 and new_text[0] in "一二三四五六七八九十" and new_text[1] == "、":
                new_text = new_text[2:]  # 删除"X、"

        elif level_type == "二级":
            # 删除"1. "或"1、"格式
            if len(new_text) >= 2 and new_text[0].isdigit() and new_text[1] in ".":
                new_text = new_text[2:]  # 删除"X."
            elif len(new_text) >= 2 and new_text[0].isdigit() and new_text[1] == "、":
                new_text = new_text[2:]  # 删除"X、"

        elif level_type == "三级":
            # 删除"（1）"格式
            if new_text.startswith("（") and len(new_text) >= 4 and new_text[2] == "）":
                new_text = new_text[4:]  # 删除"（X）"

        # 更新段落文本（简化实现，避免复杂的XML操作）
        if new_text != text:
            # 清除所有runs
            for run in para.runs:
                run._element.getparent().remove(run._element)

            # 添加新的run（简单的文本替换）
            new_run = para.add_run(new_text.strip())

    def _apply_numbering(self, para: Any, num_id: str, ilvl: str) -> None:
        """应用编号到段落"""
        pPr = para._element.get_or_add_pPr()

        # 创建 numPr 元素
        numPr = OxmlElement("w:numPr")

        # 添加 ilvl（缩进级别）
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), ilvl)
        numPr.append(ilvl_elem)

        # 添加 numId（编号ID）
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), num_id)
        numPr.append(numId_elem)

        # 添加到段落属性
        pPr.append(numPr)

    def _clear_old_format(self, pPr: Any) -> None:
        """清除旧的格式定义"""
        for tag in ["w:spacing", "w:ind", "w:jc", "w:numPr"]:
            old = pPr.find(qn(tag))
            if old is not None:
                pPr.remove(old)
