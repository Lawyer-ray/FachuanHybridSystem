from __future__ import annotations

import logging
import uuid
from pathlib import Path

from django.conf import settings
from django.core.files.uploadedfile import UploadedFile
from django.db.models import QuerySet
from docx import Document

from apps.contract_review.models.review_task import ProcessStep, ReviewTask, TaskStatus
from apps.contract_review.services.content_extractor import ContentExtractor
from apps.contract_review.services.contract_reviewer import ContractReviewer
from apps.contract_review.services.docx_formatter import DocxFormatter
from apps.contract_review.services.docx_revision_tool import DocxRevisionTool
from apps.contract_review.services.exceptions import ContractReviewError, ExtractionError
from apps.contract_review.services.heading_numbering import HeadingNumbering
from apps.contract_review.services.page_numbering import PageNumbering
from apps.contract_review.services.party_identifier import PartyIdentifier
from apps.contract_review.services.title_extractor import TitleExtractor
from apps.contract_review.services.typo_checker import TypoChecker
from apps.core.llm.service import LLMService, get_llm_service

logger = logging.getLogger(__name__)


def _upload_dir() -> Path:
    d = Path(settings.MEDIA_ROOT) / "contract_review" / "uploads"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _output_dir() -> Path:
    d = Path(settings.MEDIA_ROOT) / "contract_review" / "output"
    d.mkdir(parents=True, exist_ok=True)
    return d


class ReviewService:
    """合同审查主编排服务"""

    def upload_contract(
        self,
        file: UploadedFile,
        user: object,
        model_name: str = "",
    ) -> ReviewTask:
        """上传合同：验证 → 保存 → 提取内容 → 识别甲乙方 → 创建任务"""
        filename = file.name or "unknown.docx"
        if not filename.lower().endswith(".docx"):
            raise ContractReviewError("仅支持 .docx 格式文件")

        # 保存文件
        task_id = uuid.uuid4()
        save_path = _upload_dir() / f"{task_id}_{filename}"
        with open(save_path, "wb") as f:
            for chunk in file.chunks():
                f.write(chunk)

        # 提取内容 + 识别甲乙方
        extractor = ContentExtractor()
        identifier = PartyIdentifier()
        try:
            paragraphs = extractor.extract_paragraphs(save_path)
            parties = identifier.identify_parties(paragraphs)
        except ExtractionError as e:
            task = ReviewTask.objects.create(
                id=task_id,
                user=user,
                original_file=str(save_path),
                status=TaskStatus.EXTRACTION_FAILED,
                error_message=str(e),
                model_name=model_name,
            )
            return task

        # 提取标题
        title_extractor = TitleExtractor()
        doc = Document(str(save_path))
        title = title_extractor.extract_title(doc)

        task = ReviewTask.objects.create(
            id=task_id,
            user=user,
            original_file=str(save_path),
            contract_title=title or filename.rsplit(".", 1)[0],
            party_a=parties.get("party_a", ""),
            party_b=parties.get("party_b", ""),
            party_c=parties.get("party_c", ""),
            party_d=parties.get("party_d", ""),
            status=TaskStatus.PARTIES_IDENTIFIED,
            model_name=model_name,
        )
        return task

    def confirm_party(
        self,
        task_id: uuid.UUID,
        represented_party: str,
        user: object,
        reviewer_name: str = "",
        selected_steps: list[str] | None = None,
    ) -> ReviewTask:
        """确认代表方，提交异步审查任务"""
        task = ReviewTask.objects.get(id=task_id)
        if task.status not in (
            TaskStatus.PARTIES_IDENTIFIED,
            TaskStatus.EXTRACTION_FAILED,
        ):
            raise ContractReviewError(f"当前状态 {task.status} 不允许确认代表方")

        # 默认全选
        default_steps = ["typo_check", "format_document", "contract_review", "review_report"]
        steps = selected_steps if selected_steps else default_steps

        task.represented_party = represented_party
        task.reviewer_name = reviewer_name.strip() or "法穿AI"
        task.status = TaskStatus.CONFIRMED
        task.selected_steps = steps
        ReviewTask.objects.filter(id=task_id).update(
            represented_party=represented_party,
            reviewer_name=reviewer_name.strip() or "法穿AI",
            status=TaskStatus.CONFIRMED,
            selected_steps=steps,
        )

        # 提交异步任务
        from django_q.tasks import async_task

        async_task(
            "apps.contract_review.services.review_service.process_review",
            str(task_id),
            timeout=1800,
        )
        logger.info("已提交审查任务: %s", task_id)
        return task

    def get_task_status(self, task_id: uuid.UUID) -> ReviewTask:
        """查询任务状态"""
        return ReviewTask.objects.get(id=task_id)

    def get_result_file(self, task_id: uuid.UUID) -> Path:
        """获取结果文件路径"""
        task = ReviewTask.objects.get(id=task_id)
        if task.status != TaskStatus.COMPLETED:
            raise ContractReviewError("任务尚未完成")
        path = Path(task.output_file)
        if not path.exists():
            raise ContractReviewError("结果文件不存在")
        return path

    def get_original_file(self, task_id: uuid.UUID) -> Path:
        """获取原始上传文件路径"""
        task = ReviewTask.objects.get(id=task_id)
        path = Path(task.original_file)
        if not path.exists():
            raise ContractReviewError("原始文件不存在")
        return path


def process_review(task_id_str: str) -> None:
    """异步执行审查流水线（由 Django-Q2 调用）"""
    task_id = uuid.UUID(task_id_str)
    try:
        task = ReviewTask.objects.get(id=task_id)
    except ReviewTask.DoesNotExist:
        logger.warning("任务不存在，跳过: %s", task_id)
        return
    # 防止重复执行：已完成/已失败的任务直接跳过
    if task.status in (TaskStatus.COMPLETED, TaskStatus.FAILED):
        logger.warning("任务已终态 (%s)，跳过: %s", task.status, task_id)
        return
    task.status = TaskStatus.PROCESSING
    ReviewTask.objects.filter(id=task_id).update(status=TaskStatus.PROCESSING)

    try:
        doc = Document(task.original_file)
        extractor = ContentExtractor()
        extraction = extractor.extract_with_mapping(Path(task.original_file))
        paragraphs = extraction.paragraphs
        index_map = extraction.index_map  # filtered_idx -> doc.paragraphs idx

        llm: LLMService = get_llm_service()
        revision_tool = DocxRevisionTool()

        # 启用修订模式（Track Changes）
        revision_tool.enable_track_changes(doc)

        steps = getattr(task, "selected_steps", None) or [
            "typo_check",
            "format_document",
            "contract_review",
            "review_report",
        ]

        # Step 1: 标题提取（始终执行）
        _update_step(task, ProcessStep.TITLE_EXTRACTION)
        title_extractor = TitleExtractor()
        title = title_extractor.extract_title(doc)
        if title:
            task.contract_title = title
            ReviewTask.objects.filter(id=task.id).update(contract_title=title)

        # Step 2: 错别字检测
        if "typo_check" in steps:
            _update_step(task, ProcessStep.TYPO_CHECK)
            typo_checker = TypoChecker(llm)
            typos = typo_checker.check_typos(paragraphs, model_name=task.model_name)
            typo_applied = 0
            for typo in typos:
                applied = _apply_to_any_paragraph(
                    doc, revision_tool, typo.original, typo.corrected, author=task.reviewer_name
                )
                if applied:
                    typo_applied += 1
            logger.info("错别字修订: %d/%d 处成功", typo_applied, len(typos))

        # Step 3: 合同审查
        if "contract_review" in steps:
            _update_step(task, ProcessStep.CONTRACT_REVIEW)
            reviewer = ContractReviewer(llm)
            reviews = reviewer.review_contract(
                paragraphs,
                task.represented_party,
                task.party_a,
                task.party_b,
                model_name=task.model_name,
            )
            review_applied = 0
            for review in reviews:
                applied = _apply_to_any_paragraph(
                    doc, revision_tool, review.original, review.suggested, author=task.reviewer_name
                )
                if applied:
                    review_applied += 1
            logger.info("合同审查修订: %d/%d 处成功", review_applied, len(reviews))

        # Step 3b: 生成评估报告
        if "review_report" in steps:
            if "contract_review" not in steps:
                reviewer = ContractReviewer(llm)
            report = reviewer.generate_report(
                paragraphs,
                task.represented_party,
                task.party_a,
                task.party_b,
                model_name=task.model_name,
            )
            if report:
                task.review_report = report
                ReviewTask.objects.filter(id=task.id).update(review_report=report)
                logger.info("评估报告已生成 (%d 字)", len(report))

        # Step 4: 格式标准化
        if "format_document" in steps:
            _update_step(task, ProcessStep.FORMAT_DOCUMENT)
            DocxFormatter().format_document(doc)

            _update_step(task, ProcessStep.PAGE_NUMBERING)
            PageNumbering().standardize(doc)

            _update_step(task, ProcessStep.HEADING_NUMBERING)
            HeadingNumbering(llm).apply_numbering(doc, model_name=task.model_name)

        # 保存输出文件
        output_name = title_extractor.generate_output_filename(
            task.contract_title,
            task_id=str(task.id),
        )
        output_path = _output_dir() / output_name
        doc.save(str(output_path))

        task.output_file = str(output_path)
        task.status = TaskStatus.COMPLETED
        ReviewTask.objects.filter(id=task.id).update(
            output_file=str(output_path),
            status=TaskStatus.COMPLETED,
            current_step="",
        )
        logger.info("审查完成: %s -> %s", task_id, output_name)

    except Exception as e:
        logger.exception("审查任务失败: %s", task_id)
        ReviewTask.objects.filter(id=task_id).update(
            status=TaskStatus.FAILED,
            error_message=str(e),
        )


def _update_step(task: ReviewTask, step: str) -> None:
    task.current_step = step
    ReviewTask.objects.filter(id=task.id).update(current_step=step)


def _apply_to_any_paragraph(
    doc: Document, tool: DocxRevisionTool, original: str, replacement: str, author: str = ""
) -> bool:
    """遍历所有段落查找原文并应用修订，不依赖 LLM 返回的段落索引"""
    for para in doc.paragraphs:
        if original in para.text:
            if tool.apply_revision(para, original, replacement, author=author or None):
                return True
    return False
