#!/usr/bin/env python3
"""
创建poi-tl示例模板
"""
import sys
from pathlib import Path

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent / 'backend'))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_complaint_template():
    """创建民事起诉状模板"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('民  事  起  诉  状')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '黑体'

    # 空行
    doc.add_paragraph()

    # 原告信息
    plaintiff_para = doc.add_paragraph()
    plaintiff_para.add_run('原告：{{plaintiff.name}}，{{plaintiff.gender}}，{{plaintiff.birthDate}}出生，')
    plaintiff_para.add_run('{{plaintiff.nationality}}，住所：{{plaintiff.address}}。')

    # 代理人信息（条件渲染）
    agent_para = doc.add_paragraph()
    agent_para.add_run('{{?plaintiff.agent}}')
    agent_para.add_run('委托诉讼代理人：{{plaintiff.agent.name}}，{{plaintiff.agent.firm}}律师。')
    agent_para.add_run('{{/plaintiff.agent}}')

    # 被告信息
    defendant_para = doc.add_paragraph()
    defendant_para.add_run('被告：{{defendant.name}}，{{defendant.gender}}，{{defendant.birthDate}}出生，')
    defendant_para.add_run('{{defendant.nationality}}，住所：{{defendant.address}}。')

    # 空行
    doc.add_paragraph()

    # 诉讼请求标题
    claims_title = doc.add_paragraph()
    claims_title.add_run('诉讼请求：')

    # 诉讼请求列表（循环）
    claims_para = doc.add_paragraph()
    claims_para.add_run('{{?claims}}')
    claims_para.add_run('{{_index + 1}}. {{=#this}}')
    claims_para.add_run('{{/claims}}')

    # 空行
    doc.add_paragraph()

    # 事实与理由标题
    facts_title = doc.add_paragraph()
    facts_title.add_run('事实与理由：')

    # 事实与理由列表（循环）
    facts_para = doc.add_paragraph()
    facts_para.add_run('{{?facts}}')
    facts_para.add_run('{{=#this}}')
    facts_para.add_run('{{/facts}}')

    # 空行
    doc.add_paragraph()

    # 证据标题
    evidence_title = doc.add_paragraph()
    evidence_title.add_run('证据和证据来源：')

    # 证据列表（循环）
    evidence_para = doc.add_paragraph()
    evidence_para.add_run('{{?evidence}}')
    evidence_para.add_run('{{_index + 1}}. {{=#this}}')
    evidence_para.add_run('{{/evidence}}')

    # 空行
    doc.add_paragraph()

    # 此致
    court_para = doc.add_paragraph()
    court_para.add_run('此致')
    court_para.add_run('{{court}}')

    # 空行
    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区
    sign_para = doc.add_paragraph()
    sign_para.add_run('                                        起诉人：{{plaintiff.name}}')
    sign_para.add_run('{{submitDate}}')

    # 保存模板
    template_path = Path('/Users/huangsong21/Downloads/Coding/AI/FachuanHybridSystem/java-services/poi-service/src/main/resources/templates/complaint/civil_complaint.docx')
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    print(f'✓ 民事起诉状模板创建成功: {template_path}')

if __name__ == '__main__':
    create_complaint_template()
